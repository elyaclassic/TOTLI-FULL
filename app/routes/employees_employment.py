"""
Xodimlar — ishga qabul qilish hujjatlari (Tier C1 5-bosqich).

Manba: employees.py:219-930 (~712 qator) dan ajratib olindi.
Endpoint path'lar o'zgarishsiz.
"""
import io
from datetime import datetime, date
from typing import Optional, List
from urllib.parse import quote

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fastapi import APIRouter, Request, Depends, Form, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from sqlalchemy.orm import Session, joinedload

from app.core import templates
from app.models.database import (
    get_db,
    User,
    Employee,
    EmploymentDoc,
    Department,
    Position,
    PieceworkTask,
)
from app.deps import require_auth, require_admin

router = APIRouter(prefix="/employees", tags=["employees-employment"])


@router.get("/hiring-docs", response_class=HTMLResponse)
async def employment_docs_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul qilish hujjatlari ro'yxati"""
    docs = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .order_by(EmploymentDoc.created_at.desc())
        .all()
    )
    return templates.TemplateResponse("employees/hiring_docs_list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Ishga qabul qilish hujjatlari"
    })


@router.get("/hiring-doc/create", response_class=HTMLResponse)
async def employment_doc_create_page(
    request: Request,
    employee_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish (xodim tanlash)."""
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    emp = db.query(Employee).filter(Employee.id == employee_id).first() if employee_id else None
    if emp:
        existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).order_by(EmploymentDoc.doc_date.desc()).first()
        if existing:
            return RedirectResponse(
                url="/employees/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan. Yangi hujjat yaratib bo'lmaydi — mavjud hujjatni ko'ring yoki tahrirlang.")
                + "&existing_doc_id=" + str(existing.id),
                status_code=303,
            )
    today_str = date.today().isoformat()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("employees/hiring_doc_form.html", {
        "request": request,
        "employees": employees,
        "emp": emp,
        "today_str": today_str,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "current_user": current_user,
        "page_title": "Ishga qabul hujjati yaratish"
    })


@router.post("/hiring-doc/create")
async def employment_doc_create(
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    monthly_rest_days: int = Form(4),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati yaratish. Har bir xodim faqat bir marta."""
    if salary < 0:
        raise HTTPException(status_code=400, detail="Maosh manfiy bo'lishi mumkin emas")
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    existing = db.query(EmploymentDoc).filter(EmploymentDoc.employee_id == emp.id).first()
    if existing:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote(f"«{emp.full_name}» allaqachon ishga qabul qilingan.")
            + "&existing_doc_id=" + str(existing.id),
            status_code=303,
        )
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url="/employees/hiring-doc/create?employee_id=" + str(employee_id) + "&error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    count = db.query(EmploymentDoc).filter(EmploymentDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IQ-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc_salary = float(salary) if salary else (emp.salary or 0)
    doc_department = (department or "").strip() or (emp.department or "").strip() or None
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon", "tue", "wed", "thu", "fri", "sat", "sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"
    doc = EmploymentDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        hire_date=hire_d,
        position=(position or "").strip() or (emp.position or "").strip() or None,
        department=doc_department,
        salary=doc_salary,
        salary_type=st,
        piecework_task_ids=",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None,
        rest_days=",".join(rest_days_clean) if rest_days_clean else None,
        probation=probation_clean,
        contract_type=ct,
        contract_end_date=end_d,
        note=note or None,
        user_id=current_user.id,
        confirmed_at=datetime.now(),
    )
    db.add(doc)
    db.flush()
    emp.salary = doc_salary
    if st:
        emp.salary_type = st
    if monthly_rest_days is not None and 0 <= monthly_rest_days <= 15:
        emp.monthly_rest_days = int(monthly_rest_days)
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None
    if hire_d:
        emp.hire_date = hire_d
    if (position or "").strip():
        emp.position = (position or "").strip()
    if doc_department:
        emp.department = doc_department
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-doc/{doc.id}?created=1", status_code=303)


@router.get("/hiring-doc/{doc_id}", response_class=HTMLResponse)
async def employment_doc_view(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishga qabul hujjati ko'rish / chop etish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"
    piecework_task_names = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
            for t in tasks:
                nm = (t.name or t.code or str(t.id))
                piecework_task_names.append(nm)
    except Exception:
        piecework_task_names = []
    return templates.TemplateResponse("employees/hiring_doc.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "piecework_task_names": piecework_task_names,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number}"
    })


@router.get("/hiring-doc/{doc_id}/contract", response_class=HTMLResponse)
async def employment_doc_contract(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasi (to'liq) — namuna asosida chop etish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"
    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []
    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {
            "mon": "dushanba", "tue": "seshanba", "wed": "chorshanba",
            "thu": "payshanba", "fri": "juma", "sat": "shanba", "sun": "yakshanba",
        }
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""
    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."
    return templates.TemplateResponse("employees/labor_contract.html", {
        "request": request,
        "doc": doc,
        "display_department": display_department,
        "selected_piecework_tasks": selected_piecework_tasks,
        "company_name": company_name,
        "employer_rep_name": employer_rep_name,
        "rest_days_display": rest_days_display,
        "current_user": current_user,
        "page_title": f"Mehnat shartnomasi {doc.number}",
    })


def _build_labor_contract_docx(doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name):
    """Shartnoma matnini Word hujjati (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    h = d.add_heading("MEHNAT SHARTNOMASI", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()
    d.add_paragraph("Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()
    d.add_paragraph(
        f"{company_name} (keyingi o'rinlarda Ish beruvchi) va fuqaro {doc.employee.full_name} "
        "(keyingi o'rinlarda Xodim), mazkur mehnat shartnomasini quyidagilar haqida tuzdilar."
    )
    d.add_paragraph()
    d.add_heading("1. UMUMIY QOIDALAR", level=1)
    hire_date_str = doc.hire_date.strftime("%d.%m.%Y") if doc.hire_date else "________________"
    pos = doc.position or (doc.employee.position if doc.employee else "") or "________________"
    d.add_paragraph(
        f"1.1. Xodim {hire_date_str} sanadan boshlab {display_department} bo'limida {pos} lavozimiga ishga qabul qilinadi."
    )
    d.add_paragraph(f"1.2. Xodimning ish joyi: {display_department}.")
    if doc.contract_type == "fixed":
        end = f" ({doc.contract_end_date.strftime('%d.%m.%Y')} gacha)" if doc.contract_end_date else ""
        d.add_paragraph(f"1.3. Mazkur shartnomaning amal qilish muddati: muayyan muddatga{end}.")
    elif doc.contract_type == "task":
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: muayyan ishni bajarish davriga.")
    else:
        d.add_paragraph("1.3. Mazkur shartnomaning amal qilish muddati: nomuayyan muddatga.")
    prob = doc.probation if doc.probation else "sinovsiz"
    d.add_paragraph(f"1.4. Sinov muddati: {prob}.")
    d.add_paragraph("1.5. Xodim lavozim yo'riqnomasi va amaldagi qonunchilikka muvofiq mehnat majburiyatlarini bajaradi.")
    d.add_paragraph()
    d.add_heading("2. TOMONLARNING HUQUQ VA MAJBURIYATLARI", level=1)
    d.add_paragraph("2.1. Ish beruvchining majburiyatlari:")
    d.add_paragraph("  • Xodimga xavfsiz va samarali mehnat qilish uchun shart-sharoitlar yaratish.")
    d.add_paragraph("  • Ichki mehnat tartibi qoidalari va lavozim yo'riqnomasi bilan tanishtirish.")
    d.add_paragraph("  • Ish haqini o'z vaqtida to'lash.")
    d.add_paragraph("2.2. Xodimning majburiyatlari:")
    d.add_paragraph("  • Mehnat intizomi va ichki tartib qoidalariga rioya qilish.")
    d.add_paragraph("  • Ish beruvchining qonuniy topshiriqlarini o'z vaqtida va aniq bajarish.")
    d.add_paragraph("  • Mehnat muhofazasi va texnika xavfsizligi talablariga rioya qilish.")
    d.add_paragraph()
    d.add_heading("3. ISH VAQTI VA DAM OLISH VAQTI", level=1)
    d.add_paragraph("3.1. Ish kuni vaqti: 09:00 dan 18:00 gacha (to'liq ish kuni asosida).")
    rest = rest_days_display if rest_days_display else "shanba va yakshanba"
    d.add_paragraph(f"3.2. Dam olish kunlari: {rest}.")
    d.add_paragraph("3.3. Qonunchilikda belgilangan tartibda dam olish/bayram kunlari ishga jalb etilishi mumkin.")
    d.add_paragraph()
    d.add_heading("4. MEHNATGA HAQ TO'LASH", level=1)
    salary_type_map = {"oylik": "Oylik", "soatlik": "Soatlik", "bo'lak": "Bo'lak", "bo'lak_oylik": "Bo'lak + oylik"}
    st = salary_type_map.get(doc.salary_type, "________________")
    d.add_paragraph(f"4.1. Ish haqi turi: {st}.")
    if doc.salary_type in ("bo'lak", "bo'lak_oylik") and selected_piecework_tasks:
        d.add_paragraph("Bo'lak ishlar va stavkalar:")
        for t in selected_piecework_tasks:
            name = t.name or t.code or str(t.id)
            price = f"{t.price_per_unit:,.0f}" if t.price_per_unit is not None else "0"
            unit = t.unit_name or "birlik"
            d.add_paragraph(f"  • {name} — {price} so'm/{unit}")
    salary_val = f"{doc.salary:,.0f}" if doc.salary else "0"
    d.add_paragraph(f"4.2. Mehnat haqi miqdori: {salary_val} so'm.")
    d.add_paragraph("4.3. Ish haqi har oyda kamida ikki marta to'lanadi.")
    d.add_paragraph()
    d.add_heading("5. XIZMAT SAFARLARI", level=1)
    d.add_paragraph("5.1. Ish zaruriyatiga ko'ra Xodim xizmat safariga yuborilishi mumkin. Xarajatlar amaldagi qonunchilikka muvofiq qoplanadi.")
    d.add_paragraph()
    d.add_heading("6. MEHNAT SHARTNOMASINI BEKOR QILISH", level=1)
    d.add_paragraph("6.1. Mehnat shartnomasi O'zbekiston Respublikasi Mehnat kodeksida belgilangan tartibda bekor qilinishi mumkin.")
    d.add_paragraph()
    d.add_heading("7. MEHNAT NIZOLARI", level=1)
    d.add_paragraph("7.1. Mehnat nizolari qonun hujjatlarida belgilangan tartibda hal qilinadi.")
    d.add_paragraph()
    d.add_heading("8. TOMONLAR REKVIZITLARI VA IMZOLARI", level=1)
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph("Manzil: O'zbekiston Respublikasi, Qo'qon shahri, Jasorat ko'chasi, 52-uy")
    d.add_paragraph("STIR: 311469106")
    d.add_paragraph("Hisob raqam: 202088409071067110001")
    d.add_paragraph('Bank: "Asaka" banki Qo\'qon filiali')
    d.add_paragraph("MFO: 00873")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim:")
    d.add_paragraph(f"F.I.O: {doc.employee.full_name}")
    d.add_paragraph(f"Kodi: {doc.employee.code or '—'}")
    d.add_paragraph(f"Telefon: {doc.employee.phone or '—'}")
    d.add_paragraph("Manzil: ____________________________")
    d.add_paragraph("Pasport: ____________________________")
    d.add_paragraph("Imzo: ______________________")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@router.get("/hiring-doc/{doc_id}/contract/export-word")
async def employment_doc_contract_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Mehnat shartnomasini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee), joinedload(EmploymentDoc.user))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    display_department = (doc.department or "").strip() or None
    if not display_department and doc.employee:
        display_department = (doc.employee.department or "").strip() or None
        if not display_department and getattr(doc.employee, "department_id", None):
            dept = db.query(Department).filter(Department.id == doc.employee.department_id).first()
            if dept:
                display_department = dept.name
    if not display_department:
        display_department = "—"
    selected_piecework_tasks = []
    try:
        raw = (doc.piecework_task_ids or "").strip()
        ids = [int(x) for x in raw.split(",") if x.strip().isdigit()] if raw else []
        if ids:
            selected_piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(ids)).order_by(PieceworkTask.name).all()
    except Exception:
        selected_piecework_tasks = []
    rest_days_display = ""
    try:
        raw_rest = (doc.rest_days or "").strip()
        codes = [x for x in raw_rest.split(",") if x]
        name_map = {"mon": "dushanba", "tue": "seshanba", "wed": "chorshanba", "thu": "payshanba", "fri": "juma", "sat": "shanba", "sun": "yakshanba"}
        names = [name_map.get(c, c) for c in codes]
        if names:
            rest_days_display = ", ".join(names)
    except Exception:
        rest_days_display = ""
    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."
    buf = _build_labor_contract_docx(
        doc, display_department, selected_piecework_tasks, rest_days_display, company_name, employer_rep_name
    )
    safe_number = (doc.number or "shartnoma").replace("/", "-").replace("\\", "-")
    filename = f"Mehnat_shartnomasi_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )


@router.post("/hiring-docs/bulk-confirm")
async def employment_docs_bulk_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan ishga qabul hujjatlarini tasdiqlash"""
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    confirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and not doc.confirmed_at:
            doc.confirmed_at = datetime.now()
            confirmed += 1
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-docs?confirmed=1&count={confirmed}", status_code=303)


@router.post("/hiring-docs/bulk-cancel-confirm")
async def employment_docs_bulk_cancel_confirm(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tanlangan hujjatlarda tasdiqlashni bekor qilish"""
    form = await request.form()
    doc_ids_raw = form.getlist("doc_ids")
    try:
        doc_ids = [int(x) for x in doc_ids_raw if str(x).strip().isdigit()]
    except (ValueError, TypeError):
        doc_ids = []
    if not doc_ids:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Hech qanday hujjat tanlanmagan."), status_code=303)
    unconfirmed = 0
    for did in doc_ids:
        doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == did).first()
        if doc and doc.confirmed_at:
            doc.confirmed_at = None
            unconfirmed += 1
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-docs?unconfirmed=1&count={unconfirmed}", status_code=303)


@router.post("/hiring-doc/{doc_id}/confirm")
async def employment_doc_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tasdiqlash"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?confirmed=1", status_code=303)


@router.post("/hiring-doc/{doc_id}/cancel-confirm")
async def employment_doc_cancel_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjati tasdiqlashni bekor qilish"""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?unconfirmed=1", status_code=303)


@router.post("/hiring-doc/{doc_id}/delete")
async def employment_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini o'chirish — faqat tasdiqlanmagan."""
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni o'chirish mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303
        )
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/employees/hiring-docs?deleted=1", status_code=303)


@router.get("/hiring-doc/{doc_id}/edit", response_class=HTMLResponse)
async def employment_doc_edit_page(
    request: Request,
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini tahrirlash"""
    doc = (
        db.query(EmploymentDoc)
        .options(joinedload(EmploymentDoc.employee))
        .filter(EmploymentDoc.id == doc_id)
        .first()
    )
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    positions = db.query(Position).filter(Position.is_active == True).order_by(Position.name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    selected_piecework_ids = []
    if doc.piecework_task_ids:
        for x in (doc.piecework_task_ids or "").split(","):
            if str(x).strip().isdigit():
                selected_piecework_ids.append(int(x.strip()))
    display_department = (doc.department or "").strip() or (getattr(doc.employee, "department", "") or "").strip() or "—"
    return templates.TemplateResponse("employees/hiring_doc_edit.html", {
        "request": request,
        "doc": doc,
        "departments": departments,
        "positions": positions,
        "piecework_tasks": piecework_tasks,
        "selected_piecework_ids": selected_piecework_ids,
        "display_department": display_department,
        "current_user": current_user,
        "page_title": f"Ishga qabul {doc.number} — tahrirlash",
    })


@router.post("/hiring-doc/{doc_id}/edit")
async def employment_doc_edit_save(
    doc_id: int,
    doc_date: str = Form(...),
    hire_date: str = Form(None),
    position: str = Form(""),
    department: str = Form(""),
    salary: float = Form(0),
    salary_type: str = Form(""),
    monthly_rest_days: int = Form(4),
    piecework_task_ids: List[int] = Form([]),
    rest_days: List[str] = Form([]),
    probation: str = Form(""),
    contract_type: str = Form("indefinite"),
    contract_end_date: str = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Ishga qabul hujjatini saqlash (tahrirlash) — faqat tasdiqlanmagan."""
    if salary < 0:
        raise HTTPException(status_code=400, detail="Maosh manfiy bo'lishi mumkin emas")
    doc = db.query(EmploymentDoc).filter(EmploymentDoc.id == doc_id).first()
    if not doc:
        return RedirectResponse(url="/employees/hiring-docs?error=Hujjat topilmadi", status_code=303)
    if doc.confirmed_at:
        return RedirectResponse(
            url="/employees/hiring-docs?error=" + quote("Tasdiqlangan hujjatni tahrirlash mumkin emas. Avval «Bekor qilish» orqali tasdiqlashni bekor qiling."),
            status_code=303,
        )
    emp = db.query(Employee).filter(Employee.id == doc.employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees/hiring-docs?error=" + quote("Xodim topilmadi"), status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(url=f"/employees/hiring-doc/{doc_id}/edit?error=" + quote("Noto'g'ri sana"), status_code=303)
    hire_d = None
    if hire_date and hire_date.strip():
        try:
            hire_d = datetime.strptime(hire_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    end_d = None
    if contract_end_date and contract_end_date.strip() and (contract_type or "").strip() == "fixed":
        try:
            end_d = datetime.strptime(contract_end_date.strip(), "%Y-%m-%d").date()
        except (ValueError, TypeError):
            pass
    st = (salary_type or "").strip() or None
    if st and st not in ("oylik", "soatlik", "bo'lak", "bo'lak_oylik"):
        st = None
    task_ids = [int(x) for x in (piecework_task_ids or []) if str(x).strip().isdigit()]
    task_ids = list(dict.fromkeys(task_ids))
    rest_days_clean = [d for d in (rest_days or []) if d in ("mon", "tue", "wed", "thu", "fri", "sat", "sun")]
    probation_clean = (probation or "").strip() or None
    ct = (contract_type or "").strip() or "indefinite"
    if ct not in ("indefinite", "fixed", "task"):
        ct = "indefinite"
    doc.doc_date = doc_d
    doc.hire_date = hire_d
    doc.position = (position or "").strip() or None
    doc.department = (department or "").strip() or None
    doc.salary = float(salary or 0)
    doc.salary_type = st
    doc.piecework_task_ids = ",".join(str(x) for x in task_ids) if (st in ("bo'lak", "bo'lak_oylik") and task_ids) else None
    doc.contract_type = ct
    doc.contract_end_date = end_d
    doc.note = (note or "").strip() or None
    doc.probation = probation_clean
    doc.rest_days = ",".join(rest_days_clean) if rest_days_clean else None
    emp.salary = doc.salary
    if st:
        emp.salary_type = st
    if monthly_rest_days is not None and 0 <= monthly_rest_days <= 15:
        emp.monthly_rest_days = int(monthly_rest_days)
    if hire_d:
        emp.hire_date = hire_d
    if doc.position:
        emp.position = doc.position
    if doc.department:
        emp.department = doc.department
    if st in ("bo'lak", "bo'lak_oylik"):
        if task_ids:
            tasks = db.query(PieceworkTask).filter(PieceworkTask.id.in_(task_ids)).all()
            emp.piecework_tasks = tasks
            emp.piecework_task_id = task_ids[0]
        else:
            emp.piecework_tasks = []
            emp.piecework_task_id = None
    db.commit()
    return RedirectResponse(url=f"/employees/hiring-doc/{doc.id}?edited=1", status_code=303)
