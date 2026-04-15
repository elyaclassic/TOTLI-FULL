"""
Ma'lumotlar bo'limi — omborlar, birliklar, kategoriyalar, narx turlari, kassalar,
bo'limlar, yo'nalishlar, foydalanuvchilar, lavozimlar, hududlar, uskunalar.
"""
import io
from datetime import datetime, date as date_type
from typing import Optional
from urllib.parse import quote, urlencode
from fastapi import APIRouter, Request, Depends, Form, File, HTTPException, UploadFile, Query
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, or_, text
from sqlalchemy.exc import OperationalError
import openpyxl

from app.core import templates
from app.models.database import (
    get_db,
    User,
    Warehouse,
    Unit,
    Category,
    PriceType,
    Product,
    ProductPrice,
    ProductPriceHistory,
    CashRegister,
    Payment,
    Department,
    Direction,
    Position,
    Region,
    Machine,
    Employee,
    Partner,
    ExpenseType,
    PieceworkTask,
    ProductionGroup,
    ProductionGroupDoc,
    production_group_members,
    PasswordChangeLog,
)
from app.deps import require_auth, require_admin
from app.utils.auth import hash_password
from app.utils.db_schema import ensure_cash_opening_balance_column

router = APIRouter(prefix="/info", tags=["info"])

# /info redirect -> home router (GET /info)


# ---------- Ro'yxatlar: Bo'limlar, Omborlar, Kassalar (admin hammasini ko'radi, qo'shish/tanlash) ----------
@router.get("/themes", response_class=HTMLResponse)
async def themes_preview(request: Request, current_user: User = Depends(require_auth)):
    """Tema tanlash sahifasi."""
    return templates.TemplateResponse("info/themes.html", {
        "request": request, "current_user": current_user, "page_title": "Tema tanlash",
    })


@router.get("/rosters", response_class=HTMLResponse)
async def info_rosters(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Admin uchun: barcha bo'limlar, omborlar va kassalar ro'yxati — ko'rish, qo'shish havolalari."""
    departments = db.query(Department).order_by(Department.name).all()
    warehouses = db.query(Warehouse).order_by(Warehouse.name).all()
    cash_registers = db.query(CashRegister).order_by(CashRegister.name).all()
    return templates.TemplateResponse("info/rosters.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Bo'limlar, Omborlar, Kassalar",
        "departments": departments,
        "warehouses": warehouses,
        "cash_registers": cash_registers,
    })


# ---------- Warehouses ----------
@router.get("/warehouses", response_class=HTMLResponse)
async def info_warehouses(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    warehouses = db.query(Warehouse).all()
    departments = db.query(Department).filter(Department.is_active == True).all()
    return templates.TemplateResponse("info/warehouses.html", {
        "request": request, 
        "warehouses": warehouses, 
        "departments": departments,
        "current_user": current_user, 
        "page_title": "Omborlar"
    })


@router.post("/warehouses/add")
async def info_warehouses_add(
    request: Request,
    name: str = Form(...),
    address: str = Form(""),
    department_id: int = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing_by_name = db.query(Warehouse).filter(Warehouse.name == name).first()
    if existing_by_name:
        raise HTTPException(status_code=400, detail=f"'{name}' nomli ombor allaqachon mavjud!")
    warehouse = Warehouse(
        name=name, 
        code=None, 
        address=address, 
        department_id=department_id if department_id else None,
        is_active=True
    )
    db.add(warehouse)
    db.commit()
    return RedirectResponse(url="/info/warehouses", status_code=303)


@router.post("/warehouses/edit/{warehouse_id}")
async def info_warehouses_edit(
    warehouse_id: int,
    name: str = Form(...),
    address: str = Form(""),
    department_id: int = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    warehouse = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not warehouse:
        raise HTTPException(status_code=404, detail="Ombor topilmadi")
    existing_by_name = db.query(Warehouse).filter(
        Warehouse.name == name,
        Warehouse.id != warehouse_id,
    ).first()
    if existing_by_name:
        raise HTTPException(status_code=400, detail=f"'{name}' nomli ombor allaqachon mavjud!")
    warehouse.name = name
    warehouse.address = address
    warehouse.department_id = department_id if department_id else None
    db.commit()
    return RedirectResponse(url="/info/warehouses", status_code=303)


@router.post("/warehouses/delete/{warehouse_id}")
async def info_warehouses_delete(warehouse_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    warehouse = db.query(Warehouse).filter(Warehouse.id == warehouse_id).first()
    if not warehouse:
        raise HTTPException(status_code=404, detail="Ombor topilmadi")
    db.delete(warehouse)
    db.commit()
    return RedirectResponse(url="/info/warehouses", status_code=303)


@router.get("/warehouses/export")
async def export_warehouses(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    warehouses = db.query(Warehouse).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Warehouses"
    ws.append(["ID", "Kod", "Nomi", "Manzil"])
    for w in warehouses:
        ws.append([w.id, w.code, w.name, w.address])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=omborlar.xlsx"},
    )


@router.get("/warehouses/template")
async def template_warehouses(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Manzil"])
    ws.append(["MAIN", "Asosiy ombor", "Toshkent sh."])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=ombor_andoza.xlsx"},
    )


@router.post("/warehouses/import")
async def import_warehouses(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]:
            continue
        code, name, address = row[0], row[1], row[2]
        warehouse = db.query(Warehouse).filter(Warehouse.code == code).first()
        if not warehouse:
            warehouse = Warehouse(code=code, name=name, address=address)
            db.add(warehouse)
        else:
            warehouse.name = name
            warehouse.address = address
        db.commit()
    return RedirectResponse(url="/info/warehouses", status_code=303)


# ---------- Units ----------
@router.get("/units", response_class=HTMLResponse)
async def info_units(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    units = db.query(Unit).all()
    return templates.TemplateResponse("info/units.html", {
        "request": request, "units": units, "current_user": current_user, "page_title": "O'lchov birliklari"
    })


@router.post("/units/add")
async def info_units_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Unit).filter(Unit.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli o'lchov birligi allaqachon mavjud!")
    unit = Unit(code=code, name=name)
    db.add(unit)
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)


@router.post("/units/edit/{unit_id}")
async def info_units_edit(
    unit_id: int,
    code: str = Form(...),
    name: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="O'lchov birligi topilmadi")
    existing = db.query(Unit).filter(Unit.code == code, Unit.id != unit_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli o'lchov birligi allaqachon mavjud!")
    unit.code = code
    unit.name = name
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)


@router.post("/units/delete/{unit_id}")
async def info_units_delete(unit_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="O'lchov birligi topilmadi")
    db.delete(unit)
    db.commit()
    return RedirectResponse(url="/info/units", status_code=303)


@router.get("/units/export")
async def export_units(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    units = db.query(Unit).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Units"
    ws.append(["ID", "Kod", "Nomi"])
    for u in units:
        ws.append([u.id, u.code, u.name])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=olchov_birliklari.xlsx"},
    )


@router.get("/units/template")
async def template_units(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi"])
    ws.append(["kg", "Kilogramm"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=birlik_andoza.xlsx"},
    )


@router.post("/units/import")
async def import_units(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]:
            continue
        code, name = row[0], row[1]
        unit = db.query(Unit).filter(Unit.code == code).first()
        if not unit:
            unit = Unit(code=code, name=name)
            db.add(unit)
        else:
            unit.name = name
        db.commit()
    return RedirectResponse(url="/info/units", status_code=303)


# ---------- Categories ----------
@router.get("/categories", response_class=HTMLResponse)
async def info_categories(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    categories = db.query(Category).all()
    return templates.TemplateResponse("info/categories.html", {
        "request": request, "categories": categories, "current_user": current_user, "page_title": "Kategoriyalar"
    })


@router.post("/categories/add")
async def info_categories_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    type: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Category).filter(Category.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli kategoriya allaqachon mavjud!")
    category = Category(code=code, name=name, type=type)
    db.add(category)
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)


@router.post("/categories/edit/{category_id}")
async def info_categories_edit(
    category_id: int,
    code: str = Form(...),
    name: str = Form(...),
    type: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    category = db.query(Category).filter(Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Kategoriya topilmadi")
    existing = db.query(Category).filter(Category.code == code, Category.id != category_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli kategoriya allaqachon mavjud!")
    category.code = code
    category.name = name
    category.type = type
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)


@router.post("/categories/delete/{category_id}")
async def info_categories_delete(category_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    category = db.query(Category).filter(Category.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Kategoriya topilmadi")
    db.delete(category)
    db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)


@router.get("/categories/export")
async def export_categories(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    categories = db.query(Category).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Categories"
    ws.append(["ID", "Kod", "Nomi", "Turi"])
    for c in categories:
        ws.append([c.id, c.code, c.name, c.type])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=kategoriyalar.xlsx"},
    )


@router.get("/categories/template")
async def template_categories(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Turi"])
    ws.append(["CAT001", "Shirinliklar", "tayyor"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=kategoriya_andoza.xlsx"},
    )


@router.post("/categories/import")
async def import_categories(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]:
            continue
        code, name, type_ = row[0], row[1], row[2]
        category = db.query(Category).filter(Category.code == code).first()
        if not category:
            category = Category(code=code, name=name, type=type_)
            db.add(category)
        else:
            category.name = name
            category.type = type_
        db.commit()
    return RedirectResponse(url="/info/categories", status_code=303)


# ---------- Price types ----------
@router.get("/price-types", response_class=HTMLResponse)
async def info_price_types(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    return templates.TemplateResponse("info/price_types.html", {
        "request": request,
        "price_types": price_types,
        "current_user": current_user,
        "page_title": "Narx turlari",
    })


@router.post("/price-types/add")
async def info_price_types_add(
    name: str = Form(...),
    code: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    code = (code or "").strip() or None
    if code and db.query(PriceType).filter(PriceType.code == code).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli narx turi allaqachon mavjud!")
    pt = PriceType(name=name, code=code, is_active=True)
    db.add(pt)
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


@router.post("/price-types/edit/{price_type_id}")
async def info_price_types_edit(
    price_type_id: int,
    name: str = Form(...),
    code: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    pt = db.query(PriceType).filter(PriceType.id == price_type_id).first()
    if not pt:
        raise HTTPException(status_code=404, detail="Narx turi topilmadi")
    code = (code or "").strip() or None
    if code and db.query(PriceType).filter(PriceType.code == code, PriceType.id != price_type_id).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli narx turi allaqachon mavjud!")
    pt.name = name
    pt.code = code
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


@router.post("/price-types/delete/{price_type_id}")
async def info_price_types_delete(
    price_type_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    pt = db.query(PriceType).filter(PriceType.id == price_type_id).first()
    if not pt:
        raise HTTPException(status_code=404, detail="Narx turi topilmadi")
    db.query(ProductPrice).filter(ProductPrice.price_type_id == price_type_id).delete()
    pt.is_active = False
    db.commit()
    return RedirectResponse(url="/info/price-types", status_code=303)


# ---------- Harajat turlari (1C uslubida) ----------
@router.get("/expense-types", response_class=HTMLResponse)
async def info_expense_types(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    """Harajat turlari — ish haqqi, ishxona harajati va boshqalar (Harajatlar hujjatida tanlash uchun)."""
    types_list = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.name).all()
    return templates.TemplateResponse("info/expense_types.html", {
        "request": request,
        "expense_types": types_list,
        "current_user": current_user,
        "page_title": "Harajat turlari",
    })


@router.post("/expense-types/add")
async def info_expense_types_add(
    name: str = Form(...),
    category: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    cat = (category or "").strip() or None
    if cat and cat not in ("Ishlab chiqarish", "Ma'muriy", "Ishlab chiqarish xarajatlari", "Ma'muriy xarajatlar"):
        cat = "Boshqa"
    et = ExpenseType(name=name.strip(), category=cat or None, is_active=True)
    db.add(et)
    db.commit()
    return RedirectResponse(url="/info/expense-types?added=1", status_code=303)


@router.post("/expense-types/edit/{expense_type_id}")
async def info_expense_types_edit(
    expense_type_id: int,
    name: str = Form(...),
    category: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    et = db.query(ExpenseType).filter(ExpenseType.id == expense_type_id).first()
    if not et:
        raise HTTPException(status_code=404, detail="Harajat turi topilmadi")
    cat = (category or "").strip() or None
    if cat and cat not in ("Ishlab chiqarish", "Ma'muriy", "Ishlab chiqarish xarajatlari", "Ma'muriy xarajatlar"):
        cat = "Boshqa"
    et.name = name.strip()
    et.category = cat
    db.commit()
    return RedirectResponse(url="/info/expense-types?updated=1", status_code=303)


@router.post("/expense-types/delete/{expense_type_id}")
async def info_expense_types_delete(
    expense_type_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    et = db.query(ExpenseType).filter(ExpenseType.id == expense_type_id).first()
    if not et:
        raise HTTPException(status_code=404, detail="Harajat turi topilmadi")
    et.is_active = False
    db.commit()
    return RedirectResponse(url="/info/expense-types?deleted=1", status_code=303)


# ---------- Prices (product prices by type) ----------
@router.get("/prices", response_class=HTMLResponse)
async def info_prices(
    request: Request,
    price_type_id: Optional[int] = None,
    search: Optional[str] = None,
    type_filter: Optional[str] = None,
    price_status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    if not price_types:
        return templates.TemplateResponse("info/prices.html", {
            "request": request,
            "products": [],
            "price_types": [],
            "current_price_type_id": None,
            "product_prices_by_type": {},
            "product_tannarx": {},
            "current_user": current_user,
            "page_title": "Narxni o'rnatish",
            "filter_search": "",
            "filter_type": "",
            "filter_price_status": "all",
        })
    current_pt_id = price_type_id or (price_types[0].id if price_types else None)
    products = db.query(Product).options(joinedload(Product.unit)).filter(
        Product.is_active == True
    ).order_by(Product.name).all()
    product_prices = db.query(ProductPrice).filter(ProductPrice.price_type_id == current_pt_id).all()
    product_prices_by_type = {pp.product_id: pp.sale_price for pp in product_prices}
    product_tannarx = {p.id: float(p.purchase_price or 0) for p in products}

    # Filtrlar: mahsulot nomi, turi, sotuv narxi holati
    search_q = (search or "").strip().lower()
    type_filter_val = (type_filter or "").strip() or None
    price_status_val = (price_status or "all").strip() or "all"
    if search_q or type_filter_val or price_status_val != "all":
        filtered_products = []
        for p in products:
            if search_q and search_q not in (p.name or "").lower() and search_q not in (p.barcode or "").lower():
                continue
            if type_filter_val and (getattr(p, "type", None) or "") != type_filter_val:
                continue
            sale_val = product_prices_by_type.get(p.id)
            if sale_val is None:
                sale_val = getattr(p, "sale_price", None)
            has_sale_price = sale_val is not None and float(sale_val or 0) > 0
            if price_status_val == "set" and not has_sale_price:
                continue
            if price_status_val == "not_set" and has_sale_price:
                continue
            filtered_products.append(p)
        products = filtered_products

    return templates.TemplateResponse("info/prices.html", {
        "request": request,
        "products": products,
        "price_types": price_types,
        "current_price_type_id": current_pt_id,
        "product_prices_by_type": product_prices_by_type,
        "product_tannarx": product_tannarx,
        "current_user": current_user,
        "page_title": "Narxni o'rnatish",
        "filter_search": search or "",
        "filter_type": type_filter or "",
        "filter_price_status": price_status_val,
    })


def _next_price_history_doc_number(db: Session) -> str:
    """Narx o'zgarishi hujjati raqami: PN-YYYYMMDD-NNN"""
    from datetime import datetime
    prefix = f"PN-{datetime.now().strftime('%Y%m%d')}-"
    last = db.query(ProductPriceHistory).filter(ProductPriceHistory.doc_number.like(f"{prefix}%")).order_by(ProductPriceHistory.id.desc()).first()
    if not last or not last.doc_number:
        num = 1
    else:
        try:
            num = int(last.doc_number.rsplit("-", 1)[-1]) + 1
        except (ValueError, IndexError):
            num = 1
    return f"{prefix}{num:03d}"


@router.post("/prices/edit/{product_id}")
async def info_prices_edit(
    product_id: int,
    purchase_price: float = Form(0),
    sale_price: float = Form(0),
    price_type_id: Optional[int] = Form(None),
    redirect_search: Optional[str] = Form(None),
    redirect_type_filter: Optional[str] = Form(None),
    redirect_price_status: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    if purchase_price < 0 or sale_price < 0:
        raise HTTPException(status_code=400, detail="Narx manfiy bo'lishi mumkin emas")
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="Mahsulot topilmadi")
    old_purchase = float(product.purchase_price or 0)
    if price_type_id:
        pp = db.query(ProductPrice).filter(
            ProductPrice.product_id == product_id,
            ProductPrice.price_type_id == price_type_id,
        ).first()
        old_sale = float(pp.sale_price if pp else 0)
        product.purchase_price = purchase_price
        if pp:
            pp.sale_price = sale_price
        else:
            db.add(ProductPrice(product_id=product_id, price_type_id=price_type_id, sale_price=sale_price))
    else:
        old_sale = float(product.sale_price or 0)
        product.purchase_price = purchase_price
        product.sale_price = sale_price
    doc_number = _next_price_history_doc_number(db)
    db.add(ProductPriceHistory(
        doc_number=doc_number,
        product_id=product_id,
        price_type_id=price_type_id,
        old_purchase_price=old_purchase,
        new_purchase_price=float(purchase_price or 0),
        old_sale_price=old_sale,
        new_sale_price=float(sale_price or 0),
        changed_by_id=current_user.id,
    ))
    db.commit()
    # Saqlashdan keyin filtrlarni saqlab qolish (faqat "Filtrni tozalash" bosilganda tozalash)
    params = {}
    if price_type_id is not None:
        params["price_type_id"] = price_type_id
    if redirect_search and str(redirect_search).strip():
        params["search"] = redirect_search.strip()
    if redirect_type_filter and str(redirect_type_filter).strip():
        params["type_filter"] = redirect_type_filter.strip()
    if redirect_price_status and str(redirect_price_status).strip() and redirect_price_status != "all":
        params["price_status"] = redirect_price_status.strip()
    redirect_url = "/info/prices" + ("?" + urlencode(params) if params else "")
    return RedirectResponse(url=redirect_url, status_code=303)


@router.get("/prices/history", response_class=HTMLResponse)
async def info_prices_history(
    request: Request,
    product_id: Optional[str] = Query(None, description="Mahsulot ID (bo'sh = barchasi)"),
    price_type_id: Optional[str] = Query(None, description="Narx turi ID (bo'sh = barchasi)"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Narx o'zgarishlari tarixi — hujjatlar ro'yxati (avvalgi va yangi narx)."""
    from sqlalchemy.orm import joinedload
    pid = None
    ptid = None
    if product_id and str(product_id).strip():
        try:
            pid = int(product_id)
        except (ValueError, TypeError):
            pass
    if price_type_id and str(price_type_id).strip():
        try:
            ptid = int(price_type_id)
        except (ValueError, TypeError):
            pass
    q = db.query(ProductPriceHistory).options(
        joinedload(ProductPriceHistory.product),
        joinedload(ProductPriceHistory.price_type),
        joinedload(ProductPriceHistory.changed_by),
    ).order_by(ProductPriceHistory.changed_at.desc())
    if pid is not None:
        q = q.filter(ProductPriceHistory.product_id == pid)
    if ptid is not None:
        q = q.filter(ProductPriceHistory.price_type_id == ptid)
    history = q.limit(500).all()
    price_types = db.query(PriceType).filter(PriceType.is_active == True).order_by(PriceType.name).all()
    products = db.query(Product).filter(Product.is_active == True).order_by(Product.name).all()
    return templates.TemplateResponse("info/price_history.html", {
        "request": request,
        "history": history,
        "price_types": price_types,
        "products": products,
        "current_user": current_user,
        "page_title": "Narx o'zgarishlari tarixi",
        "show_tannarx": (getattr(current_user, "role", None) if current_user else None) == "admin",
    })


# ---------- Cash ----------
def _cash_balance_formula(db: Session, cash_id: int) -> float:
    """Kassa balansi = qoldiq (opening_balance) + kirim - chiqim (faqat tasdiqlangan)."""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        return 0.0
    opening = float(getattr(cash, "opening_balance", None) or 0)
    confirmed = or_(Payment.status == "confirmed", Payment.status.is_(None))
    income_sum = float(
        db.query(func.coalesce(func.sum(Payment.amount), 0))
        .filter(Payment.cash_register_id == cash_id, Payment.type == "income", confirmed)
        .scalar()
    ) or 0
    expense_sum = float(
        db.query(func.coalesce(func.sum(Payment.amount), 0))
        .filter(Payment.cash_register_id == cash_id, Payment.type == "expense", confirmed)
        .scalar()
    ) or 0
    return opening + income_sum - expense_sum


@router.get("/cash", response_class=HTMLResponse)
async def info_cash(request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    ensure_cash_opening_balance_column(db)
    cash_registers = db.query(CashRegister).all()
    departments = db.query(Department).filter(Department.is_active == True).all()
    cash_computed_balance = {}
    for c in cash_registers:
        cash_computed_balance[c.id] = _cash_balance_formula(db, c.id)
    jami_balans = sum(cash_computed_balance.values())
    return templates.TemplateResponse("info/cash.html", {
        "request": request,
        "cash_registers": cash_registers,
        "cash_computed_balance": cash_computed_balance,
        "jami_balans": jami_balans,
        "departments": departments,
        "current_user": current_user,
        "page_title": "Kassalar",
    })


@router.post("/cash/add")
async def info_cash_add(
    request: Request,
    name: str = Form(...),
    balance: float = Form(0),
    department_id: int = Form(None),
    payment_type: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    ensure_cash_opening_balance_column(db)
    pt = (payment_type or "").strip() or None
    if pt and pt not in ("naqd", "plastik", "click", "terminal"):
        pt = None
    cash = CashRegister(
        name=name,
        balance=float(balance),
        opening_balance=float(balance),
        department_id=department_id if department_id else None,
        payment_type=pt,
        is_active=True,
    )
    db.add(cash)
    db.commit()
    db.refresh(cash)
    cash.balance = _cash_balance_formula(db, cash.id)
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)


@router.post("/cash/edit/{cash_id}")
async def info_cash_edit(
    cash_id: int,
    name: str = Form(...),
    balance: float = Form(0),
    department_id: int = Form(None),
    payment_type: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    ensure_cash_opening_balance_column(db)
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    cash.name = name
    cash.opening_balance = float(balance)
    cash.department_id = department_id if department_id else None
    pt = (payment_type or "").strip() or None
    cash.payment_type = pt if pt in ("naqd", "plastik", "click", "terminal") else None
    cash.balance = _cash_balance_formula(db, cash_id)
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)


@router.post("/cash/delete/{cash_id}")
async def info_cash_delete(cash_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    db.delete(cash)
    db.commit()
    return RedirectResponse(url="/info/cash", status_code=303)


@router.post("/cash/recalculate/{cash_id}")
async def info_cash_recalculate(
    cash_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Kassa balansini formuladan qayta yozadi: balans = qoldiq (opening_balance) + kirim - chiqim (faqat admin)."""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    ensure_cash_opening_balance_column(db)
    cash.balance = _cash_balance_formula(db, cash_id)
    db.commit()
    return RedirectResponse(
        url="/info/cash?recalculated=1&balance=" + quote(str(cash.balance)),
        status_code=303,
    )


# ---------- Departments ----------
@router.get("/departments", response_class=HTMLResponse)
async def info_departments(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    departments = db.query(Department).all()
    return templates.TemplateResponse("info/departments.html", {
        "request": request,
        "departments": departments,
        "current_user": current_user,
        "page_title": "Bo'limlar",
    })


@router.post("/departments/add")
async def info_departments_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Department).filter(Department.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lim allaqachon mavjud!")
    department = Department(code=code, name=name, description=description, is_active=True)
    db.add(department)
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)


@router.post("/departments/edit/{department_id}")
async def info_departments_edit(
    department_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Bo'lim topilmadi")
    existing = db.query(Department).filter(Department.code == code, Department.id != department_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lim allaqachon mavjud!")
    department.code = code
    department.name = name
    department.description = description
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)


@router.post("/departments/delete/{department_id}")
async def info_departments_delete(department_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    department = db.query(Department).filter(Department.id == department_id).first()
    if not department:
        raise HTTPException(status_code=404, detail="Bo'lim topilmadi")
    db.delete(department)
    db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)


@router.get("/departments/export")
async def export_departments(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    departments = db.query(Department).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Departments"
    ws.append(["ID", "Kod", "Nomi", "Izoh"])
    for d in departments:
        ws.append([d.id, d.code, d.name, d.description])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=bolimlar.xlsx"},
    )


@router.get("/departments/template")
async def template_departments(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Izoh"])
    ws.append(["DEP001", "Ishlab chiqarish", "Asosiy tsex"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=bolim_andoza.xlsx"},
    )


@router.post("/departments/import")
async def import_departments(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]:
            continue
        code, name, description = row[0], row[1], row[2]
        department = db.query(Department).filter(Department.code == code).first()
        if not department:
            department = Department(code=code, name=name, description=description)
            db.add(department)
        else:
            department.name = name
            department.description = description
        db.commit()
    return RedirectResponse(url="/info/departments", status_code=303)


# ---------- Directions ----------
@router.get("/directions", response_class=HTMLResponse)
async def info_directions(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    directions = db.query(Direction).all()
    return templates.TemplateResponse("info/directions.html", {
        "request": request,
        "directions": directions,
        "current_user": current_user,
        "page_title": "Yo'nalishlar",
    })


@router.post("/directions/add")
async def info_directions_add(
    request: Request,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Direction).filter(Direction.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli yo'nalish allaqachon mavjud!")
    direction = Direction(code=code, name=name, description=description, is_active=True)
    db.add(direction)
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)


@router.post("/directions/edit/{direction_id}")
async def info_directions_edit(
    direction_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    direction = db.query(Direction).filter(Direction.id == direction_id).first()
    if not direction:
        raise HTTPException(status_code=404, detail="Yo'nalish topilmadi")
    existing = db.query(Direction).filter(Direction.code == code, Direction.id != direction_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli yo'nalish allaqachon mavjud!")
    direction.code = code
    direction.name = name
    direction.description = description
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)


@router.post("/directions/delete/{direction_id}")
async def info_directions_delete(direction_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    direction = db.query(Direction).filter(Direction.id == direction_id).first()
    if not direction:
        raise HTTPException(status_code=404, detail="Yo'nalish topilmadi")
    db.delete(direction)
    db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)


@router.get("/directions/export")
async def export_directions(db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    directions = db.query(Direction).all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Directions"
    ws.append(["ID", "Kod", "Nomi", "Izoh"])
    for d in directions:
        ws.append([d.id, d.code, d.name, d.description])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=yonalishlar.xlsx"},
    )


@router.get("/directions/template")
async def template_directions(current_user: User = Depends(require_auth)):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Template"
    ws.append(["Kod", "Nomi", "Izoh"])
    ws.append(["DIR001", "Halva", "Halva mahsulotlari"])
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=yonalish_andoza.xlsx"},
    )


@router.post("/directions/import")
async def import_directions(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    contents = await file.read()
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for row in rows:
        if not row[0]:
            continue
        code, name, description = row[0], row[1], row[2]
        direction = db.query(Direction).filter(Direction.code == code).first()
        if not direction:
            direction = Direction(code=code, name=name, description=description)
            db.add(direction)
        else:
            direction.name = name
            direction.description = description
        db.commit()
    return RedirectResponse(url="/info/directions", status_code=303)


# ---------- Users (admin) ----------
@router.get("/users", response_class=HTMLResponse)
async def info_users(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_admin)
):
    users = (
        db.query(User)
        .options(
            joinedload(User.department),
            joinedload(User.warehouse),
            joinedload(User.cash_register),
            joinedload(User.departments_list),
            joinedload(User.warehouses_list),
            joinedload(User.cash_registers_list),
            joinedload(User.partners_list),
        )
        .order_by(User.id)
        .all()
    )
    # Ishga qabul qilingan (faol) hodimlar — Xodim dropdown uchun
    employees = (
        db.query(Employee)
        .filter(Employee.is_active == True)
        .order_by(Employee.full_name)
        .all()
    )
    user_to_employee = {e.user_id: e for e in db.query(Employee).filter(Employee.user_id != None).all()}
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    warehouses = db.query(Warehouse).filter(Warehouse.is_active == True).order_by(Warehouse.name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    error = request.query_params.get("error", "").strip()
    return templates.TemplateResponse("info/users.html", {
        "request": request,
        "users": users,
        "employees": employees,
        "user_to_employee": user_to_employee,
        "departments": departments,
        "warehouses": warehouses,
        "cash_registers": cash_registers,
        "partners": partners,
        "current_user": current_user,
        "page_title": "Foydalanuvchilar",
        "error": error,
    })


def _parse_id_list(form_list) -> list:
    out = []
    for x in form_list:
        try:
            v = int(x)
            if v > 0 and v not in out:
                out.append(v)
        except (ValueError, TypeError):
            pass
    return out


@router.post("/users/add")
async def info_users_add(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    full_name: str = Form(...),
    role: str = Form("user"),
    is_active: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    form = await request.form()
    department_ids = _parse_id_list(form.getlist("department_ids"))
    warehouse_ids = _parse_id_list(form.getlist("warehouse_ids"))
    cash_register_ids = _parse_id_list(form.getlist("cash_register_ids"))
    partner_ids = _parse_id_list(form.getlist("partner_ids"))
    existing = db.query(User).filter(User.username == username).first()
    if existing:
        msg = quote(f"'{username}' login bilan foydalanuvchi allaqachon mavjud! Boshqa login tanlang.")
        return RedirectResponse(url=f"/info/users?error={msg}", status_code=303)
    user = User(
        username=username,
        password_hash=hash_password(password),
        full_name=full_name,
        role=role,
        is_active=is_active,
        department_id=department_ids[0] if department_ids else None,
        warehouse_id=warehouse_ids[0] if warehouse_ids else None,
        cash_register_id=cash_register_ids[0] if cash_register_ids else None,
    )
    db.add(user)
    db.flush()
    for did in department_ids:
        dept = db.query(Department).filter(Department.id == did).first()
        if dept:
            user.departments_list.append(dept)
    for wid in warehouse_ids:
        wh = db.query(Warehouse).filter(Warehouse.id == wid).first()
        if wh:
            user.warehouses_list.append(wh)
    for cid in cash_register_ids:
        cash = db.query(CashRegister).filter(CashRegister.id == cid).first()
        if cash:
            user.cash_registers_list.append(cash)
    for pid in partner_ids:
        partner = db.query(Partner).filter(Partner.id == pid).first()
        if partner:
            user.partners_list.append(partner)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)


@router.post("/users/edit/{user_id}")
async def info_users_edit(
    user_id: int,
    request: Request,
    username: str = Form(...),
    full_name: str = Form(...),
    role: str = Form("user"),
    is_active: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    form = await request.form()
    department_ids = _parse_id_list(form.getlist("department_ids"))
    warehouse_ids = _parse_id_list(form.getlist("warehouse_ids"))
    cash_register_ids = _parse_id_list(form.getlist("cash_register_ids"))
    partner_ids = _parse_id_list(form.getlist("partner_ids"))
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    existing = db.query(User).filter(User.username == username, User.id != user_id).first()
    if existing:
        msg = quote(f"'{username}' login bilan foydalanuvchi allaqachon mavjud! Boshqa login tanlang.")
        return RedirectResponse(url=f"/info/users?error={msg}", status_code=303)
    user.username = username
    user.full_name = full_name
    user.role = role
    user.is_active = is_active
    user.department_id = department_ids[0] if department_ids else None
    user.warehouse_id = warehouse_ids[0] if warehouse_ids else None
    user.cash_register_id = cash_register_ids[0] if cash_register_ids else None
    user.departments_list.clear()
    user.warehouses_list.clear()
    user.cash_registers_list.clear()
    user.partners_list.clear()
    for did in department_ids:
        dept = db.query(Department).filter(Department.id == did).first()
        if dept:
            user.departments_list.append(dept)
    for wid in warehouse_ids:
        wh = db.query(Warehouse).filter(Warehouse.id == wid).first()
        if wh:
            user.warehouses_list.append(wh)
    for cid in cash_register_ids:
        cash = db.query(CashRegister).filter(CashRegister.id == cid).first()
        if cash:
            user.cash_registers_list.append(cash)
    for pid in partner_ids:
        partner = db.query(Partner).filter(Partner.id == pid).first()
        if partner:
            user.partners_list.append(partner)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)


@router.post("/users/change-password/{user_id}")
async def info_users_change_password(
    request: Request,
    user_id: int,
    new_password: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    user.password_hash = hash_password(new_password)

    # Audit log
    client_ip = request.client.host if request.client else None
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        client_ip = forwarded.split(",")[0].strip()
    user_agent = (request.headers.get("User-Agent") or "")[:255]

    log = PasswordChangeLog(
        target_user_id=user.id,
        target_username=user.username,
        changed_by_id=current_user.id,
        changed_by_username=current_user.username,
        ip_address=client_ip,
        user_agent=user_agent,
    )
    db.add(log)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)


@router.get("/users/password-logs", response_class=HTMLResponse)
async def info_users_password_logs(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Parol o'zgartirish audit jurnali (faqat admin)."""
    logs = (
        db.query(PasswordChangeLog)
        .order_by(PasswordChangeLog.changed_at.desc())
        .limit(500)
        .all()
    )
    return templates.TemplateResponse("info/password_logs.html", {
        "request": request,
        "current_user": current_user,
        "logs": logs,
        "page_title": "Parol o'zgartirishlar jurnali",
    })


@router.post("/users/delete/{user_id}")
async def info_users_delete(
    user_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_admin)
):
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="O'zingizni o'chira olmaysiz!")
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Foydalanuvchi topilmadi")
    db.delete(user)
    db.commit()
    return RedirectResponse(url="/info/users", status_code=303)


# ---------- Positions ----------
@router.get("/positions", response_class=HTMLResponse)
async def info_positions(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    positions = db.query(Position).filter(Position.is_active == True).all()
    return templates.TemplateResponse("info/positions.html", {
        "request": request,
        "current_user": current_user,
        "positions": positions,
        "page_title": "Lavozimlar",
    })


@router.post("/positions/add")
async def info_positions_add(
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Position).filter(Position.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli lavozim allaqachon mavjud!")
    position = Position(code=code, name=name, description=description or None)
    db.add(position)
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


@router.post("/positions/edit/{position_id}")
async def info_positions_edit(
    position_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    position = db.query(Position).filter(Position.id == position_id).first()
    if not position:
        raise HTTPException(status_code=404, detail="Lavozim topilmadi")
    existing = db.query(Position).filter(Position.code == code, Position.id != position_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli lavozim allaqachon mavjud!")
    position.code = code
    position.name = name
    position.description = description or None
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


@router.post("/positions/delete/{position_id}")
async def info_positions_delete(
    position_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    position = db.query(Position).filter(Position.id == position_id).first()
    if not position:
        raise HTTPException(status_code=404, detail="Lavozim topilmadi")
    position.is_active = False
    db.commit()
    return RedirectResponse(url="/info/positions", status_code=303)


# ---------- Piecework tasks (bo'lak ishlar) ----------
@router.get("/piecework-tasks", response_class=HTMLResponse)
async def info_piecework_tasks(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bo'lak ishlar (bo'lak turlari) ro'yxati — ishga qabul va oylikda ishlatiladi."""
    tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    return templates.TemplateResponse("info/piecework_tasks.html", {
        "request": request,
        "current_user": current_user,
        "tasks": tasks,
        "page_title": "Bo'lak ishlar",
    })


@router.post("/piecework-tasks/add", response_class=RedirectResponse)
async def info_piecework_tasks_add(
    code: str = Form(""),
    name: str = Form(""),
    price_per_unit: float = Form(0),
    unit_name: str = Form("kg"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bo'lak ish qo'shish"""
    if price_per_unit < 0:
        raise HTTPException(status_code=400, detail="Narx manfiy bo'lishi mumkin emas")
    code = (code or "").strip() or None
    if code and db.query(PieceworkTask).filter(PieceworkTask.code == code).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lak ish allaqachon mavjud!")
    task = PieceworkTask(
        code=code,
        name=(name or "").strip() or None,
        price_per_unit=float(price_per_unit or 0),
        unit_name=(unit_name or "kg").strip() or "kg",
    )
    db.add(task)
    db.commit()
    return RedirectResponse(url="/info/piecework-tasks?added=1", status_code=303)


@router.post("/piecework-tasks/edit/{task_id}", response_class=RedirectResponse)
async def info_piecework_tasks_edit(
    task_id: int,
    code: str = Form(""),
    name: str = Form(""),
    price_per_unit: float = Form(0),
    unit_name: str = Form("kg"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Bo'lak ishni tahrirlash"""
    if price_per_unit < 0:
        raise HTTPException(status_code=400, detail="Narx manfiy bo'lishi mumkin emas")
    task = db.query(PieceworkTask).filter(PieceworkTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Bo'lak ish topilmadi")
    code = (code or "").strip() or None
    if code and db.query(PieceworkTask).filter(PieceworkTask.code == code, PieceworkTask.id != task_id).first():
        raise HTTPException(status_code=400, detail=f"'{code}' kodli bo'lak ish allaqachon mavjud!")
    task.code = code
    task.name = (name or "").strip() or None
    task.price_per_unit = float(price_per_unit or 0)
    task.unit_name = (unit_name or "kg").strip() or "kg"
    db.commit()
    return RedirectResponse(url="/info/piecework-tasks?updated=1", status_code=303)


@router.post("/piecework-tasks/delete/{task_id}", response_class=RedirectResponse)
async def info_piecework_tasks_delete(
    task_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    """Bo'lak ishni o'chirish (soft - is_active=False)"""
    task = db.query(PieceworkTask).filter(PieceworkTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Bo'lak ish topilmadi")
    task.is_active = False
    db.commit()
    return RedirectResponse(url="/info/piecework-tasks?deleted=1", status_code=303)


# ---------- Production groups (ishlab chiqarish guruhlari) ----------
@router.get("/production-groups", response_class=HTMLResponse)
async def info_production_groups(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarish guruhlari — operator + a'zolar, kunlik tabel bo'yicha bo'lak taqsimlanadi."""
    groups = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.operator), joinedload(ProductionGroup.piecework_task), joinedload(ProductionGroup.members))
        .filter(ProductionGroup.is_active == True)
        .order_by(ProductionGroup.name)
        .all()
    )
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    # Har bir guruh uchun oxirgi hujjat
    last_docs = {}
    if groups:
        group_ids = [g.id for g in groups]
        from sqlalchemy import desc
        all_docs = (
            db.query(ProductionGroupDoc)
            .filter(ProductionGroupDoc.group_id.in_(group_ids))
            .order_by(ProductionGroupDoc.id.desc())
            .all()
        )
        for d in all_docs:
            if d.group_id not in last_docs:
                last_docs[d.group_id] = d
    return templates.TemplateResponse("info/production_groups.html", {
        "request": request,
        "groups": groups,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "last_docs": last_docs,
        "current_user": current_user,
        "page_title": "Ishlab chiqarish guruhlari (qiyomchilar)",
    })


@router.post("/production-groups/add", response_class=RedirectResponse)
async def info_production_groups_add(
    request: Request,
    name: str = Form(...),
    operator_id: int = Form(...),
    piecework_task_id: Optional[int] = Form(None),
    include_qiyom: str = Form("1"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruh qo'shish + avtomatik buyruq hujjati yaratish."""
    from datetime import datetime, date as date_type
    form = await request.form()
    member_ids_raw = form.getlist("member_ids") if hasattr(form, "getlist") else []
    member_ids = list(dict.fromkeys([int(x) for x in member_ids_raw if str(x).strip().isdigit()]))
    gr = ProductionGroup(
        name=name.strip(),
        operator_id=operator_id,
        piecework_task_id=int(piecework_task_id) if piecework_task_id else None,
        include_qiyom=(include_qiyom == "1"),
    )
    db.add(gr)
    db.flush()
    for eid in member_ids:
        db.execute(production_group_members.insert().values(group_id=gr.id, employee_id=eid))
    # Buyruq hujjati yaratish
    doc = _create_production_group_doc(db, gr, member_ids, "create", current_user)
    db.commit()
    return RedirectResponse(url=f"/info/production-group-doc/{doc.id}?created=1", status_code=303)


@router.get("/production-groups/edit/{group_id}", response_class=HTMLResponse)
async def info_production_groups_edit_page(
    group_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruhni tahrirlash sahifasi."""
    gr = (
        db.query(ProductionGroup)
        .options(joinedload(ProductionGroup.operator), joinedload(ProductionGroup.piecework_task), joinedload(ProductionGroup.members))
        .filter(ProductionGroup.id == group_id)
        .first()
    )
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    employees = db.query(Employee).filter(Employee.is_active == True).order_by(Employee.full_name).all()
    piecework_tasks = db.query(PieceworkTask).filter(PieceworkTask.is_active == True).order_by(PieceworkTask.name).all()
    # Har bir a'zo uchun narx
    member_prices = {}
    rows = db.execute(
        production_group_members.select().where(production_group_members.c.group_id == group_id)
    ).fetchall()
    for row in rows:
        member_prices[row.employee_id] = row.price_per_unit or 0
    return templates.TemplateResponse("info/production_group_edit.html", {
        "request": request,
        "group": gr,
        "employees": employees,
        "piecework_tasks": piecework_tasks,
        "member_prices": member_prices,
        "current_user": current_user,
        "page_title": f"Guruhni tahrirlash: {gr.name}",
    })


@router.post("/production-groups/edit/{group_id}", response_class=RedirectResponse)
async def info_production_groups_edit(
    group_id: int,
    request: Request,
    name: str = Form(...),
    operator_id: int = Form(...),
    piecework_task_id: Optional[int] = Form(None),
    include_qiyom: str = Form("1"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Guruhni saqlash + yangi buyruq hujjati yaratish."""
    from sqlalchemy import delete
    from datetime import datetime, date as date_type
    form = await request.form()
    member_ids_raw = form.getlist("member_ids") if hasattr(form, "getlist") else []
    member_ids = list(dict.fromkeys([int(x) for x in member_ids_raw if str(x).strip().isdigit()]))
    gr = db.query(ProductionGroup).filter(ProductionGroup.id == group_id).first()
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    gr.name = name.strip()
    gr.operator_id = operator_id
    gr.piecework_task_id = int(piecework_task_id) if piecework_task_id else None
    gr.include_qiyom = (include_qiyom == "1")
    # Yaratilgan sanani o'zgartirish
    created_date_raw = form.get("created_date", "")
    if created_date_raw and created_date_raw.strip():
        try:
            gr.created_at = datetime.strptime(created_date_raw.strip(), "%Y-%m-%d")
        except (ValueError, TypeError):
            pass
    db.execute(delete(production_group_members).where(production_group_members.c.group_id == group_id))
    for eid in member_ids:
        price_raw = form.get(f"price_{eid}", "0")
        try:
            price = float(price_raw) if price_raw else 0
        except (ValueError, TypeError):
            price = 0
        db.execute(production_group_members.insert().values(group_id=gr.id, employee_id=eid, price_per_unit=price))
    # Yangi buyruq hujjati (o'zgartirish)
    doc = _create_production_group_doc(db, gr, member_ids, "update", current_user)
    db.commit()
    return RedirectResponse(url=f"/info/production-group-doc/{doc.id}?created=1", status_code=303)


@router.post("/production-groups/delete/{group_id}", response_class=RedirectResponse)
async def info_production_groups_delete(
    group_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    """Guruhni o'chirish (is_active=False)."""
    gr = db.query(ProductionGroup).filter(ProductionGroup.id == group_id).first()
    if not gr:
        raise HTTPException(status_code=404, detail="Guruh topilmadi")
    gr.is_active = False
    db.commit()
    return RedirectResponse(url="/info/production-groups?deleted=1", status_code=303)


def _build_production_group_docx(doc):
    """Ishlab chiqarish guruhi buyrug'ini Word (.docx) sifatida qaytaradi (BytesIO)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = DocxDocument()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    title = "ISHLAB CHIQARISH GURUHINI TASHKIL ETISH HAQIDA BUYRUQ" if doc.doc_type == "create" else "ISHLAB CHIQARISH GURUHINI O'ZGARTIRISH HAQIDA BUYRUQ"
    d.add_heading(title, level=0)
    h = d.paragraphs[-1]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()
    d.add_paragraph(f"1. Guruh nomi: {doc.group_name or '—'}")
    d.add_paragraph(f"2. Operator: {doc.operator_name or '—'}")
    d.add_paragraph(f"3. Bo'lak narxi: {doc.piecework_info or '—'}")
    d.add_paragraph(f"4. Qiyom kiritilsin: {'Ha' if doc.include_qiyom else 'Yo`q'}")
    d.add_paragraph()
    d.add_paragraph("5. Guruh a'zolari:")
    if doc.members_snapshot:
        for i, line in enumerate(doc.members_snapshot.split("; "), 1):
            d.add_paragraph(f"   {i}. {line}")
    else:
        d.add_paragraph("   — a'zo yo'q")
    if doc.note:
        d.add_paragraph()
        d.add_paragraph(f"Izoh: {doc.note}")
    d.add_paragraph()
    d.add_paragraph(f"Kuchga kirish sanasi: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    d.add_paragraph()
    d.add_paragraph("Rahbar: ______________________")
    d.add_paragraph("Imzo: ______________________")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


import os

DOCS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "docs")
os.makedirs(DOCS_DIR, exist_ok=True)


def _create_production_group_doc(db, gr, member_ids, doc_type, current_user):
    """Ishlab chiqarish guruhi uchun buyruq hujjati yaratadi (snapshot bilan) va .docx faylga saqlaydi."""
    from app.models.database import Employee, PieceworkTask
    today = date_type.today()
    count = db.query(ProductionGroupDoc).filter(ProductionGroupDoc.doc_date >= today.replace(day=1)).count()
    number = f"IGR-{today.strftime('%Y%m%d')}-{count + 1:04d}"
    # Operator nomi
    operator = db.query(Employee).filter(Employee.id == gr.operator_id).first()
    operator_name = operator.full_name if operator else str(gr.operator_id)
    # Bo'lak narxi
    piecework_info = "—"
    if gr.piecework_task_id:
        pt = db.query(PieceworkTask).filter(PieceworkTask.id == gr.piecework_task_id).first()
        if pt:
            piecework_info = f"{pt.name or pt.code} — {pt.price_per_unit:,.0f} so'm/{pt.unit_name or 'kg'}"
    # A'zolar snapshot
    members_lines = []
    if member_ids:
        emps = db.query(Employee).filter(Employee.id.in_(member_ids)).order_by(Employee.full_name).all()
        # Narxlar
        rate_rows = db.execute(
            production_group_members.select().where(production_group_members.c.group_id == gr.id)
        ).fetchall()
        rates = {r.employee_id: r.price_per_unit for r in rate_rows}
        for e in emps:
            price = rates.get(e.id, 0) or 0
            members_lines.append(f"{e.full_name} ({e.code or '—'}) — {price:,.0f} so'm" if price else f"{e.full_name} ({e.code or '—'})")
    members_snapshot = "; ".join(members_lines) if members_lines else "—"
    # Kuchga kirish sanasi: tashkil etish = guruh sanasi, o'zgartirish = bugun
    if doc_type == "create":
        effective_date = gr.created_at.date() if gr.created_at else today
    else:
        effective_date = today
    doc = ProductionGroupDoc(
        number=number,
        group_id=gr.id,
        doc_date=effective_date,
        doc_type=doc_type,
        group_name=gr.name,
        operator_name=operator_name,
        piecework_info=piecework_info,
        include_qiyom=gr.include_qiyom,
        members_snapshot=members_snapshot,
        user_id=current_user.id if current_user else None,
    )
    db.add(doc)
    db.flush()
    # .docx faylni serverga saqlash
    buf = _build_production_group_docx(doc)
    safe_name = number.replace("/", "_")
    file_path = os.path.join(DOCS_DIR, f"{safe_name}.docx")
    with open(file_path, "wb") as f:
        f.write(buf.read())
    return doc


@router.get("/production-group-docs", response_class=HTMLResponse)
async def production_group_docs_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarish guruhi buyruqlari ro'yxati."""
    docs = (
        db.query(ProductionGroupDoc)
        .options(joinedload(ProductionGroupDoc.group), joinedload(ProductionGroupDoc.user))
        .order_by(ProductionGroupDoc.id.desc())
        .all()
    )
    return templates.TemplateResponse("info/production_group_docs_list.html", {
        "request": request,
        "docs": docs,
        "current_user": current_user,
        "page_title": "Ishlab chiqarish guruhi buyruqlari",
    })


@router.get("/production-group-doc/{doc_id}", response_class=HTMLResponse)
async def production_group_doc_view(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarish guruhi buyrug'i — ko'rish."""
    doc = (
        db.query(ProductionGroupDoc)
        .options(joinedload(ProductionGroupDoc.group), joinedload(ProductionGroupDoc.user))
        .filter(ProductionGroupDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    # Shu guruhning barcha hujjatlari
    all_docs = (
        db.query(ProductionGroupDoc)
        .filter(ProductionGroupDoc.group_id == doc.group_id)
        .order_by(ProductionGroupDoc.doc_date.desc(), ProductionGroupDoc.id.desc())
        .all()
    )
    return templates.TemplateResponse("info/production_group_doc.html", {
        "request": request,
        "doc": doc,
        "all_docs": all_docs,
        "current_user": current_user,
        "page_title": f"Guruh buyrug'i {doc.number}",
    })


@router.get("/production-group-doc/{doc_id}/export-word")
async def production_group_doc_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishlab chiqarish guruhi buyrug'ini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(ProductionGroupDoc)
        .options(joinedload(ProductionGroupDoc.group))
        .filter(ProductionGroupDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    safe_name = (doc.number or "doc").replace("/", "_")
    file_path = os.path.join(DOCS_DIR, f"{safe_name}.docx")
    # Agar fayl mavjud bo'lsa — uni qaytarish, aks holda generatsiya
    if not os.path.exists(file_path):
        buf = _build_production_group_docx(doc)
        with open(file_path, "wb") as f:
            f.write(buf.read())
    return StreamingResponse(
        open(file_path, "rb"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )


@router.post("/production-group-doc/{doc_id}/update-date", response_class=RedirectResponse)
async def production_group_doc_update_date(
    doc_id: int,
    doc_date: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Buyruqning kuchga kirish sanasini o'zgartirish (faqat qoralama holatda)."""
    doc = db.query(ProductionGroupDoc).filter(ProductionGroupDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "draft":
        raise HTTPException(status_code=400, detail="Faqat qoralama hujjat sanasini o'zgartirish mumkin")
    try:
        doc.doc_date = datetime.strptime(doc_date.strip(), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        pass
    db.commit()
    return RedirectResponse(url=f"/info/production-group-doc/{doc_id}", status_code=303)


@router.post("/production-group-doc/{doc_id}/confirm", response_class=RedirectResponse)
async def production_group_doc_confirm(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Buyruqni tasdiqlash."""
    doc = db.query(ProductionGroupDoc).filter(ProductionGroupDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    doc.status = "confirmed"
    doc.confirmed_at = datetime.now()
    db.commit()
    return RedirectResponse(url=f"/info/production-group-doc/{doc_id}?confirmed=1", status_code=303)


@router.post("/production-group-doc/{doc_id}/cancel", response_class=RedirectResponse)
async def production_group_doc_cancel(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Buyruqni bekor qilish."""
    doc = db.query(ProductionGroupDoc).filter(ProductionGroupDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    doc.status = "cancelled"
    doc.confirmed_at = None
    db.commit()
    return RedirectResponse(url=f"/info/production-group-doc/{doc_id}?cancelled=1", status_code=303)


@router.post("/production-group-doc/{doc_id}/delete", response_class=RedirectResponse)
async def production_group_doc_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Buyruqni o'chirish."""
    doc = db.query(ProductionGroupDoc).filter(ProductionGroupDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    # Faylni ham o'chirish
    safe_name = (doc.number or "doc").replace("/", "_")
    file_path = os.path.join(DOCS_DIR, f"{safe_name}.docx")
    if os.path.exists(file_path):
        os.remove(file_path)
    db.delete(doc)
    db.commit()
    return RedirectResponse(url="/info/production-group-docs?deleted=1", status_code=303)


# ---------- Regions ----------
@router.get("/regions", response_class=HTMLResponse)
async def info_regions(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    regions = db.query(Region).filter(Region.is_active == True).all()
    return templates.TemplateResponse("info/regions.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Hududlar",
        "regions": regions,
    })


@router.post("/regions/add")
async def region_add(
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    existing = db.query(Region).filter(Region.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli hudud allaqachon mavjud!")
    region = Region(code=code, name=name, description=description)
    db.add(region)
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


@router.post("/regions/edit/{region_id}")
async def region_edit(
    region_id: int,
    code: str = Form(...),
    name: str = Form(...),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    region = db.query(Region).filter(Region.id == region_id).first()
    if not region:
        raise HTTPException(status_code=404, detail="Hudud topilmadi")
    existing = db.query(Region).filter(Region.code == code, Region.id != region_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli hudud allaqachon mavjud!")
    region.code = code
    region.name = name
    region.description = description
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


@router.post("/regions/delete/{region_id}")
async def region_delete(region_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)):
    region = db.query(Region).filter(Region.id == region_id).first()
    if not region:
        raise HTTPException(status_code=404, detail="Hudud topilmadi")
    db.delete(region)
    db.commit()
    return RedirectResponse(url="/info/regions", status_code=303)


# ---------- Machines ----------
@router.get("/machines", response_class=HTMLResponse)
async def info_machines(
    request: Request, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machines = db.query(Machine).filter(Machine.is_active == True).order_by(Machine.created_at.desc()).all()
    employees = db.query(Employee).filter(Employee.is_active == True).all()
    warehouses = db.query(Warehouse).all()
    return templates.TemplateResponse("info/machines.html", {
        "request": request,
        "current_user": current_user,
        "page_title": "Uskunalar",
        "machines": machines,
        "employees": employees,
        "warehouses": warehouses,
    })


@router.post("/machines/add")
async def machine_add(
    code: str = Form(...),
    name: str = Form(...),
    machine_type: str = Form(""),
    capacity: float = Form(0),
    efficiency: float = Form(100.0),
    status: str = Form("idle"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    existing = db.query(Machine).filter(Machine.code == code).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli uskuna allaqachon mavjud!")
    machine = Machine(
        code=code.strip(),
        name=name.strip(),
        machine_type=machine_type.strip() or "boshqa",
        capacity=float(capacity),
        efficiency=float(efficiency),
        status=status,
    )
    db.add(machine)
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)


@router.post("/machines/edit/{machine_id}")
async def machine_edit(
    machine_id: int,
    code: str = Form(...),
    name: str = Form(...),
    machine_type: str = Form(""),
    capacity: float = Form(0),
    efficiency: float = Form(100.0),
    status: str = Form("idle"),
    operator_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machine = db.query(Machine).filter(Machine.id == machine_id).first()
    if not machine:
        raise HTTPException(status_code=404, detail="Uskuna topilmadi")
    existing = db.query(Machine).filter(Machine.code == code, Machine.id != machine_id).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"'{code}' kodli uskuna allaqachon mavjud!")
    machine.code = code.strip()
    machine.name = name.strip()
    machine.machine_type = machine_type.strip() or "boshqa"
    machine.capacity = float(capacity)
    machine.efficiency = float(efficiency)
    machine.status = status
    machine.operator_id = int(operator_id) if operator_id else None
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)


@router.post("/machines/delete/{machine_id}")
async def machine_delete(
    machine_id: int, db: Session = Depends(get_db), current_user: User = Depends(require_auth)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    machine = db.query(Machine).filter(Machine.id == machine_id).first()
    if not machine:
        raise HTTPException(status_code=404, detail="Uskuna topilmadi")
    machine.is_active = False
    db.commit()
    return RedirectResponse(url="/info/machines", status_code=303)
