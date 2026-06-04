"""
Moliya — kassa, to'lovlar, harajatlar, harajat turlari, kassadan kassaga o'tkazish.
"""
from datetime import datetime, timedelta
from typing import Optional
from urllib.parse import quote

from fastapi import APIRouter, Request, Depends, Form, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_, func, text
from sqlalchemy.exc import OperationalError, IntegrityError

from app.core import templates
from app.models.database import (
    get_db, User, CashRegister, Payment, CashTransfer,
    Partner, Purchase, PurchaseExpense, ExpenseDoc, ExpenseDocItem, ExpenseType,
    Direction, Department, EmployeeAdvance, Employee, Salary,
)
import re as _re
from app.deps import require_auth, require_admin
from app.services.period_service import is_period_closed
from app.utils.db_schema import ensure_payments_status_column, ensure_cash_opening_balance_column
from app.utils.audit import log_action
from app.utils.doc_number import next_doc_number
from app.constants import QUERY_LIMIT_DEFAULT

router = APIRouter(prefix="/finance", tags=["finance"])
cash_router = APIRouter(prefix="/cash", tags=["cash-transfers"])


def _cash_balance_formula(db: Session, cash_id: int) -> tuple:
    from app.services.finance_service import cash_balance_formula
    return cash_balance_formula(db, cash_id)


def _sync_cash_balance(db: Session, cash_id: int) -> None:
    """Wrapper — haqiqiy logika finance_service.sync_cash_balance da."""
    from app.services.finance_service import sync_cash_balance
    sync_cash_balance(db, cash_id)


def _next_expense_doc_number(db: Session) -> str:
    # H5 fix: id.desc() emas, prefiks bo'yicha MAX(suffix)+1 (uniform helper)
    today = datetime.now().strftime("%Y%m%d")
    return next_doc_number(db, ExpenseDoc, f"HD-{today}-")


def _transfer_conversion(db: Session, from_cash_id: int, to_cash_id: int, amount: float):
    """M8: o'tkazma uchun (exchange_rate, to_amount) qaytaradi.

    Cross-currency (from/to valyutasi farqli) bo'lsa joriy kurs (currency_service.get_rate)
    bilan to_amount = amount*rate; bir xil valyuta bo'lsa (None, None). Kurs topilmasa
    ValueError("FROM->TO"). Edit'da to_amount/rate eskirib qolmasligi uchun ishlatiladi.
    """
    fc = db.query(CashRegister).filter(CashRegister.id == from_cash_id).first()
    tc = db.query(CashRegister).filter(CashRegister.id == to_cash_id).first()
    from_curr = (getattr(fc, "currency", None) or "UZS").upper() if fc else "UZS"
    to_curr = (getattr(tc, "currency", None) or "UZS").upper() if tc else "UZS"
    if from_curr == to_curr:
        return None, None
    from app.services.currency_service import get_rate
    rate = get_rate(db, from_curr, to_curr)
    if not rate or rate <= 0:
        raise ValueError(f"{from_curr}->{to_curr}")
    return float(rate), round(float(amount) * float(rate), 2)


@router.get("", response_class=HTMLResponse)
async def finance(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    cash_register_id: Optional[str] = None,
    partner_id: Optional[str] = None,
):
    """Moliya - kassa. So'nggi to'lovlar sana, kassa va kontragent bo'yicha filtrlanishi mumkin."""
    ensure_payments_status_column(db)
    role = (current_user.role or "").strip()
    # Sotuvchi uchun cheklov: faqat biriktirilgan kassalar va mijozlar
    user_scope = db.query(User).options(
        joinedload(User.cash_registers_list),
        joinedload(User.partners_list),
    ).filter(User.id == current_user.id).first()
    allowed_cash_ids = None
    allowed_partner_ids = None
    if role == "sotuvchi":
        allowed_cash_ids = [c.id for c in (getattr(user_scope, "cash_registers_list", None) or []) if c]
        allowed_partner_ids = [p.id for p in (getattr(user_scope, "partners_list", None) or []) if p]
        cash_registers = db.query(CashRegister).filter(
            CashRegister.id.in_(allowed_cash_ids or [-1])
        ).order_by(CashRegister.name).all()
        if allowed_partner_ids:
            partners = db.query(Partner).filter(
                Partner.is_active == True,
                Partner.id.in_(allowed_partner_ids),
            ).order_by(Partner.name).all()
        else:
            partners = []
    else:
        cash_registers = db.query(CashRegister).all()
        partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    q = (
        db.query(Payment)
        .options(joinedload(Payment.cash_register), joinedload(Payment.partner))
        .order_by(Payment.date.desc())
    )
    if role == "sotuvchi":
        q = q.filter(Payment.cash_register_id.in_(allowed_cash_ids or [-1]))
    # Sana filtrlari
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    has_date_filter = bool(filter_date_from or filter_date_to)
    df_parsed = None
    dt_parsed = None
    if filter_date_from:
        try:
            df_parsed = datetime.strptime(filter_date_from, "%Y-%m-%d").date()
            q = q.filter(Payment.date >= df_parsed)
        except ValueError:
            pass
    if filter_date_to:
        try:
            dt_parsed = datetime.strptime(filter_date_to, "%Y-%m-%d").date()
            q = q.filter(Payment.date < datetime.combine(dt_parsed + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    # Kassa va kontragent filtrlari
    filter_cash_id = None
    filter_partner_id = None
    try:
        if cash_register_id and str(cash_register_id).strip():
            filter_cash_id = int(cash_register_id)
            q = q.filter(Payment.cash_register_id == filter_cash_id)
    except (ValueError, TypeError):
        filter_cash_id = None
    try:
        if partner_id and str(partner_id).strip():
            filter_partner_id = int(partner_id)
            q = q.filter(Payment.partner_id == filter_partner_id)
    except (ValueError, TypeError):
        filter_partner_id = None
    from app.utils.pagination import paginate, pagination_query_string
    _pg = paginate(q, request.query_params.get("page", 1), per_page=50)
    payments = _pg["items"]
    # Har payment uchun ExpenseDoc va EmployeeAdvance linklari (klikab havolalar uchun)
    payment_ids_all = [p.id for p in payments if p.id]
    expense_doc_by_payment = {}
    advance_by_payment = {}
    if payment_ids_all:
        for ed in db.query(ExpenseDoc.id, ExpenseDoc.payment_id, ExpenseDoc.number).filter(
            ExpenseDoc.payment_id.in_(payment_ids_all)
        ).all():
            expense_doc_by_payment[ed.payment_id] = {"id": ed.id, "number": ed.number}
        # Avans linklari
        avans_payments = [p for p in payments if p.description and p.description.startswith("Avans:")]
        if avans_payments:
            p_dates = [p.date for p in avans_payments if p.date]
            if p_dates:
                advs = db.query(EmployeeAdvance).options(joinedload(EmployeeAdvance.employee)).filter(
                    EmployeeAdvance.confirmed_at.isnot(None),
                    EmployeeAdvance.advance_date >= min(p_dates).date(),
                    EmployeeAdvance.advance_date <= max(p_dates).date(),
                ).all()
                adv_idx = {}
                for a in advs:
                    key = (round(float(a.amount or 0), 2), a.advance_date, a.cash_register_id)
                    adv_idx.setdefault(key, []).append(a)
                used_ids = set()
                for p in avans_payments:
                    if not p.date: continue
                    key = (round(float(p.amount or 0), 2), p.date.date(), p.cash_register_id)
                    for a in adv_idx.get(key, []):
                        if a.id in used_ids: continue
                        emp_name = (a.employee.full_name or "")[:100] if a.employee else ""
                        if emp_name and emp_name in p.description:
                            advance_by_payment[p.id] = a.id
                            used_ids.add(a.id)
                            break
    today = datetime.now().date()
    _status_ok = or_(Payment.status == "confirmed", Payment.status.is_(None))

    # Stat kartochkalar — sana filtri bo'lsa shu oraliq, bo'lmasa bugungi
    if has_date_filter:
        stat_q_income = db.query(Payment).filter(Payment.type == "income", _status_ok)
        stat_q_expense = db.query(Payment).filter(Payment.type == "expense", _status_ok)
        if df_parsed:
            stat_q_income = stat_q_income.filter(Payment.date >= df_parsed)
            stat_q_expense = stat_q_expense.filter(Payment.date >= df_parsed)
        if dt_parsed:
            stat_q_income = stat_q_income.filter(Payment.date < datetime.combine(dt_parsed + timedelta(days=1), datetime.min.time()))
            stat_q_expense = stat_q_expense.filter(Payment.date < datetime.combine(dt_parsed + timedelta(days=1), datetime.min.time()))
        # P10 audit fix: SQL func.sum() (avval Python sum() barcha qatorlarni xotiraga olardi)
        stat_income = float(stat_q_income.with_entities(func.coalesce(func.sum(Payment.amount), 0)).scalar() or 0)
        stat_expense = float(stat_q_expense.with_entities(func.coalesce(func.sum(Payment.amount), 0)).scalar() or 0)
        stats_label = f"{filter_date_from} — {filter_date_to}" if filter_date_from and filter_date_to else (filter_date_from or filter_date_to)
    else:
        # P10 audit fix: SQL func.sum() — DB tomonida agregatsiya
        try:
            stat_income = float(db.query(func.coalesce(func.sum(Payment.amount), 0)).filter(
                Payment.type == "income", Payment.date >= today, _status_ok
            ).scalar() or 0)
            stat_expense = float(db.query(func.coalesce(func.sum(Payment.amount), 0)).filter(
                Payment.type == "expense", Payment.date >= today, _status_ok
            ).scalar() or 0)
        except OperationalError:
            stat_income = float(db.query(func.coalesce(func.sum(Payment.amount), 0)).filter(
                Payment.type == "income", Payment.date >= today
            ).scalar() or 0)
            stat_expense = float(db.query(func.coalesce(func.sum(Payment.amount), 0)).filter(
                Payment.type == "expense", Payment.date >= today
            ).scalar() or 0)
        stats_label = "Bugungi"
    stats = {
        "income": stat_income,
        "expense": stat_expense,
        "label": stats_label,
    }

    # Kassalar — sana filtri bo'lsa shu oraliq bo'yicha kirim/chiqim (bitta GROUP BY query)
    cash_data = []
    if has_date_filter:
        cq = db.query(
            Payment.cash_register_id,
            Payment.type,
            func.coalesce(func.sum(Payment.amount), 0).label("total"),
        ).filter(_status_ok)
        if df_parsed:
            cq = cq.filter(Payment.date >= df_parsed)
        if dt_parsed:
            cq = cq.filter(Payment.date < datetime.combine(dt_parsed + timedelta(days=1), datetime.min.time()))
        cq = cq.group_by(Payment.cash_register_id, Payment.type).all()
        stats_map: dict = {}
        for row in cq:
            stats_map.setdefault(row.cash_register_id, {"income": 0.0, "expense": 0.0})
            stats_map[row.cash_register_id][row.type] = float(row.total or 0)
        as_of = dt_parsed or df_parsed
        for cash in cash_registers:
            s = stats_map.get(cash.id, {"income": 0.0, "expense": 0.0})
            bal_at_end, _, _ = _cash_balance_formula_at(db, cash.id, as_of)
            cash_data.append({
                "cash": cash,
                "balance": bal_at_end,
                "income": s["income"],
                "expense": s["expense"],
                "is_filtered": True,
            })
    else:
        for cash in cash_registers:
            cash_data.append({
                "cash": cash,
                "balance": float(cash.balance or 0),
                "income": 0,
                "expense": 0,
                "is_filtered": False,
            })

    return templates.TemplateResponse("finance/index.html", {
        "request": request,
        "cash_registers": cash_registers,
        "cash_data": cash_data,
        "partners": partners,
        "payments": payments,
        "expense_doc_by_payment": expense_doc_by_payment,
        "advance_by_payment": advance_by_payment,
        "stats": stats,
        "has_date_filter": has_date_filter,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "filter_cash_id": filter_cash_id,
        "filter_partner_id": filter_partner_id,
        "current_user": current_user,
        "page": _pg["page"],
        "per_page": _pg["per_page"],
        "total_count": _pg["total_count"],
        "total_pages": _pg["total_pages"],
        "items_count": _pg["items_count"],
        "base_url": "/finance",
        "pagination_query": pagination_query_string({
            "date_from": filter_date_from,
            "date_to": filter_date_to,
            "cash_register_id": filter_cash_id or "",
            "partner_id": filter_partner_id or "",
        }),
        "page_title": "Moliya"
    })


@router.get("/harajatlar", response_class=HTMLResponse)
async def finance_harajatlar(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    show_all: Optional[str] = None,
    cash_id: Optional[int] = None,
    direction_id: Optional[int] = None,
    department_id: Optional[int] = None,
    kind: Optional[str] = None,
):
    """Harajatlar jurnali — harajat hujjatlari va boshqa chiqimlar (1C uslubida).

    Default: faqat bugun. ?show_all=1 → barcha sanalar. ?date_from=..&date_to=.. → aniq oraliq.
    Filtrlar: cash_id (kassa), direction_id (yo'nalish), department_id (bo'lim),
    kind (tur: hd|xarid|other|expense|sale_return). Yo'nalish/Bo'lim HD hujjatlarga (va
    xaridlarga) tegishli — ular o'rnatilsa yo'nalishsiz oddiy to'lovlar chiqarib tashlanadi.
    """
    ensure_payments_status_column(db)
    kind = (kind or "").strip() or None
    # Yo'nalish/Bo'lim filtri o'rnatilsa, yo'nalishsiz oddiy Payment'lar chiqariladi
    _dirdept = bool(direction_id or department_id)
    if date_from is None and date_to is None and show_all != "1":
        _today_str = datetime.now().date().isoformat()
        date_from = _today_str
        date_to = _today_str
    cash_registers = db.query(CashRegister).all()
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    directions = db.query(Direction).filter(Direction.is_active == True).order_by(Direction.name).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    expense_docs_q = (
        db.query(ExpenseDoc)
        .filter(ExpenseDoc.status != "deleted")
        .options(
            joinedload(ExpenseDoc.cash_register),
            joinedload(ExpenseDoc.direction),
            joinedload(ExpenseDoc.department),
        )
    )
    # Filtrlar (HD hujjatlar): kassa, yo'nalish, bo'lim
    if cash_id:
        expense_docs_q = expense_docs_q.filter(ExpenseDoc.cash_register_id == cash_id)
    if direction_id:
        expense_docs_q = expense_docs_q.filter(ExpenseDoc.direction_id == direction_id)
    if department_id:
        expense_docs_q = expense_docs_q.filter(ExpenseDoc.department_id == department_id)
    # Tur filtri: HD'ni faqat kind bo'sh yoki 'expense_doc' bo'lganda ko'rsatamiz
    _show_hd = kind in (None, "expense_doc")
    if not _show_hd:
        expense_docs_q = expense_docs_q.filter(ExpenseDoc.id == -1)
    _df_main = str(date_from or "").strip()[:10] if date_from else ""
    _dt_main = str(date_to or "").strip()[:10] if date_to else ""
    if _df_main:
        try:
            df_x = datetime.strptime(_df_main, "%Y-%m-%d").date()
            expense_docs_q = expense_docs_q.filter(ExpenseDoc.date >= datetime.combine(df_x, datetime.min.time()))
        except ValueError:
            pass
    if _dt_main:
        try:
            dt_x = datetime.strptime(_dt_main, "%Y-%m-%d").date()
            expense_docs_q = expense_docs_q.filter(ExpenseDoc.date < datetime.combine(dt_x + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    expense_docs = expense_docs_q.order_by(ExpenseDoc.date.desc()).limit(100).all()
    purchases_with_expenses_q = (
        db.query(Purchase)
        .options(
            joinedload(Purchase.expense_cash_register),
            joinedload(Purchase.expense_direction),
            joinedload(Purchase.expense_department),
        )
        .filter(Purchase.status == "confirmed", Purchase.total_expenses > 0)
    )
    # Filtrlar (xarid xarajatlari): kassa/yo'nalish/bo'lim (expense_* ustunlari)
    if cash_id:
        purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.expense_cash_register_id == cash_id)
    if direction_id:
        purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.expense_direction_id == direction_id)
    if department_id:
        purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.expense_department_id == department_id)
    _show_xarid = kind in (None, "supplier_payment")
    if not _show_xarid:
        purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.id == -1)
    _df = str(date_from or "").strip()[:10] if date_from else ""
    _dt = str(date_to or "").strip()[:10] if date_to else ""
    if _df:
        try:
            df = datetime.strptime(_df, "%Y-%m-%d").date()
            purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.date >= datetime.combine(df, datetime.min.time()))
        except ValueError:
            pass
    if _dt:
        try:
            dt = datetime.strptime(_dt, "%Y-%m-%d").date()
            purchases_with_expenses_q = purchases_with_expenses_q.filter(Purchase.date < datetime.combine(dt + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    purchases_with_expenses = purchases_with_expenses_q.order_by(Purchase.date.desc()).limit(100).all()
    harajat_hujjatlari = []
    for doc in expense_docs:
        harajat_hujjatlari.append({
            "is_purchase_expense_doc": False,
            "date": doc.date,
            "cash_register": doc.cash_register,
            "direction": doc.direction,
            "department": doc.department,
            "total_amount": doc.total_amount or 0,
            "status": doc.status or "draft",
            "url": f"/finance/harajat/hujjat/{doc.id}",
            "number": doc.number or "",
            "doc_id": doc.id,
        })
    for p in purchases_with_expenses:
        harajat_hujjatlari.append({
            "is_purchase_expense_doc": True,
            "date": p.date,
            "cash_register": p.expense_cash_register,
            "direction": getattr(p, "expense_direction", None),
            "department": getattr(p, "expense_department", None),
            "total_amount": p.total_expenses or 0,
            "status": "confirmed",
            "url": f"/purchases/edit/{p.id}",
            "number": p.number or "",
            "purchase_id": p.id,
        })
    harajat_hujjatlari.sort(key=lambda x: x["date"] or datetime.min, reverse=True)
    harajat_hujjatlari = harajat_hujjatlari[:150]
    q = (
        db.query(Payment)
        .options(joinedload(Payment.cash_register), joinedload(Payment.partner))
        .filter(Payment.type == "expense", Payment.status != "cancelled")
        .order_by(Payment.date.desc())
    )
    # Filtrlar (PAY/boshqa chiqim): kassa, kategoriya; yo'nalish/bo'lim o'rnatilsa
    # yo'nalishsiz oddiy to'lovlar chiqariladi (HD/xarid yuqorida alohida ko'rsatiladi)
    if cash_id:
        q = q.filter(Payment.cash_register_id == cash_id)
    if kind:
        q = q.filter(Payment.category == kind)
    if _dirdept:
        q = q.filter(Payment.id == -1)
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date >= df)
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date < datetime.combine(dt + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    payments = q.limit(QUERY_LIMIT_DEFAULT).all()
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    purchase_expenses_q = (
        db.query(PurchaseExpense)
        .options(
            joinedload(PurchaseExpense.purchase).joinedload(Purchase.partner),
            joinedload(PurchaseExpense.purchase).joinedload(Purchase.expense_cash_register),
        )
        .join(Purchase, PurchaseExpense.purchase_id == Purchase.id)
        .filter(Purchase.status == "confirmed")
    )
    # Filtrlar (xarid xarajat satrlari): expense_* ustunlari
    if cash_id:
        purchase_expenses_q = purchase_expenses_q.filter(Purchase.expense_cash_register_id == cash_id)
    if direction_id:
        purchase_expenses_q = purchase_expenses_q.filter(Purchase.expense_direction_id == direction_id)
    if department_id:
        purchase_expenses_q = purchase_expenses_q.filter(Purchase.expense_department_id == department_id)
    if kind and kind != "supplier_payment":
        purchase_expenses_q = purchase_expenses_q.filter(Purchase.id == -1)
    if filter_date_from:
        try:
            df = datetime.strptime(filter_date_from[:10], "%Y-%m-%d").date()
            purchase_expenses_q = purchase_expenses_q.filter(Purchase.date >= datetime.combine(df, datetime.min.time()))
        except ValueError:
            pass
    if filter_date_to:
        try:
            dt = datetime.strptime(filter_date_to[:10], "%Y-%m-%d").date()
            purchase_expenses_q = purchase_expenses_q.filter(Purchase.date < datetime.combine(dt + timedelta(days=1), datetime.min.time()))
        except ValueError:
            pass
    purchase_expenses_list = purchase_expenses_q.order_by(Purchase.date.desc()).limit(QUERY_LIMIT_DEFAULT).all()
    # Har payment uchun bog'langan ExpenseDoc (HD-... raqamlari klikab bo'lishi uchun)
    payment_ids_expense = [p.id for p in payments if p.id]
    expense_doc_by_payment = {}
    if payment_ids_expense:
        for ed in db.query(ExpenseDoc.id, ExpenseDoc.payment_id, ExpenseDoc.number).filter(
            ExpenseDoc.payment_id.in_(payment_ids_expense)
        ).all():
            expense_doc_by_payment[ed.payment_id] = {"id": ed.id, "number": ed.number}
    # "Avans: ..." bilan boshlanadigan payment lar uchun EmployeeAdvance ni topish
    # (amount + date + cash_register mosligi bilan)
    advance_by_payment = {}
    avans_payments = [p for p in payments if p.description and p.description.startswith("Avans:")]
    if avans_payments:
        # Barcha mos keluvchi avanslar (sana oralig'i paymentlar bilan mos)
        p_dates = [p.date for p in avans_payments if p.date]
        if p_dates:
            advs = db.query(EmployeeAdvance).options(joinedload(EmployeeAdvance.employee)).filter(
                EmployeeAdvance.confirmed_at.isnot(None),
                EmployeeAdvance.advance_date >= min(p_dates).date(),
                EmployeeAdvance.advance_date <= max(p_dates).date(),
            ).all()
            # Index: (amount, advance_date, cash_register_id) -> list of advances
            adv_idx = {}
            for a in advs:
                key = (round(float(a.amount or 0), 2), a.advance_date, a.cash_register_id)
                adv_idx.setdefault(key, []).append(a)
            used_ids = set()
            for p in avans_payments:
                if not p.date: continue
                key = (round(float(p.amount or 0), 2), p.date.date(), p.cash_register_id)
                candidates = adv_idx.get(key, [])
                # First unused match
                for a in candidates:
                    if a.id in used_ids: continue
                    # Name check
                    emp_name = (a.employee.full_name or "")[:100] if a.employee else ""
                    if emp_name and emp_name in p.description:
                        advance_by_payment[p.id] = a.id
                        used_ids.add(a.id)
                        break
    all_outflows = []
    for p in payments:
        all_outflows.append({
            "is_purchase_expense": False,
            "date": p.date,
            "amount": float(p.amount or 0),
            "description": p.description or "-",
            "partner": p.partner,
            "cash_register": p.cash_register,
            "payment": p,
            "purchase_id": None,
            "expense_doc": expense_doc_by_payment.get(p.id),
            "advance_id": advance_by_payment.get(p.id),
        })
    for pe in purchase_expenses_list:
        pu = pe.purchase
        all_outflows.append({
            "is_purchase_expense": True,
            "date": pu.date if pu else datetime.now(),
            "amount": float(pe.amount or 0),
            "description": f"Kirim xarajati: {pu.number or ''} — {pe.name or 'xarajat'}",
            "partner": pu.partner if pu else None,
            "cash_register": pu.expense_cash_register if pu else None,
            "payment": None,
            "purchase_id": pu.id if pu else None,
            "expense_doc": None,
            "advance_id": None,
        })
    all_outflows.sort(key=lambda x: x["date"] or datetime.min, reverse=True)
    all_outflows = all_outflows[:200]
    # Stat oralig'i: filter bo'lsa filter, aks holda bugun
    if _df_main:
        try:
            stat_from = datetime.strptime(_df_main, "%Y-%m-%d").date()
        except ValueError:
            stat_from = datetime.now().date()
    else:
        stat_from = datetime.now().date()
    if _dt_main:
        try:
            stat_to_excl = datetime.strptime(_dt_main, "%Y-%m-%d").date() + timedelta(days=1)
        except ValueError:
            stat_to_excl = stat_from + timedelta(days=1)
    else:
        stat_to_excl = stat_from + timedelta(days=1)

    # audit_correction — kassa balansini moslash, real harajat emas (istisno qilinadi)
    def _apply_stat_filters(pq):
        # Stat kartalari filtrlarga mos bo'lsin (kassa/tur/yo'nalish/bo'lim)
        if cash_id:
            pq = pq.filter(Payment.cash_register_id == cash_id)
        if kind:
            pq = pq.filter(Payment.category == kind)
        if _dirdept:
            pq = pq.join(ExpenseDoc, ExpenseDoc.payment_id == Payment.id)
            if direction_id:
                pq = pq.filter(ExpenseDoc.direction_id == direction_id)
            if department_id:
                pq = pq.filter(ExpenseDoc.department_id == department_id)
        return pq
    try:
        _status_ok = or_(Payment.status == "confirmed", Payment.status.is_(None))
        period_payments = _apply_stat_filters(db.query(Payment).filter(
            Payment.type == "expense",
            Payment.date >= stat_from,
            Payment.date < stat_to_excl,
            or_(Payment.category != "audit_correction", Payment.category.is_(None)),
            _status_ok
        )).all()
    except OperationalError:
        period_payments = _apply_stat_filters(db.query(Payment).filter(
            Payment.type == "expense",
            Payment.date >= stat_from,
            Payment.date < stat_to_excl,
        )).all()
    # Payment'lardan qaysilari ExpenseDoc bilan bog'langan (HD-...)
    payment_ids = [p.id for p in period_payments if p.id]
    hd_payment_ids = set()
    if payment_ids:
        for row in db.query(ExpenseDoc.payment_id).filter(
            ExpenseDoc.payment_id.in_(payment_ids),
            ExpenseDoc.status != "deleted",
        ).all():
            if row[0]:
                hd_payment_ids.add(row[0])
    docs_sum = sum(float(p.amount or 0) for p in period_payments if p.id in hd_payment_ids)
    other_sum = sum(float(p.amount or 0) for p in period_payments if p.id not in hd_payment_ids)
    stats = {
        "today_income": 0,
        "today_expense": docs_sum + other_sum,
        "today_expense_docs": docs_sum,
        "today_expense_other": other_sum,
    }
    return templates.TemplateResponse("finance/harajatlar.html", {
        "request": request,
        "cash_registers": cash_registers,
        "partners": partners,
        "expense_docs": expense_docs,
        "harajat_hujjatlari": harajat_hujjatlari,
        "payments": payments,
        "all_outflows": all_outflows,
        "stats": stats,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "directions": directions,
        "departments": departments,
        "sel_cash_id": cash_id,
        "sel_direction_id": direction_id,
        "sel_department_id": department_id,
        "sel_kind": kind or "",
        "kind_options": [
            ("expense_doc", "Harajat hujjati (HD)"),
            ("supplier_payment", "Xarid xarajati"),
            ("other", "Boshqa to'lov"),
            ("expense", "Oddiy harajat"),
            ("sale_return", "Sotuv qaytarish"),
        ],
        "current_user": current_user,
        "page_title": "Harajatlar jurnali",
        "finance_harajatlar": True,
    })


@router.get("/kassa/{cash_register_id}", response_class=HTMLResponse)
async def finance_kassa_detail(
    cash_register_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    page: Optional[int] = None,
):
    """Kassaning kirim/chiqimlari."""
    cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash:
        raise HTTPException(status_code=404, detail="Kassa topilmadi")
    ensure_cash_opening_balance_column(db)
    computed_balance, total_income_all_time, total_expense_all_time = _cash_balance_formula(db, cash_register_id)
    total_income_all_time = float(total_income_all_time)
    total_expense_all_time = float(total_expense_all_time)
    stored_balance = float(cash.balance or 0)
    balance_mismatch = abs(computed_balance - stored_balance) > 0.01
    # ---------- Payments ----------
    q = (
        db.query(Payment)
        .options(joinedload(Payment.partner))
        .filter(Payment.cash_register_id == cash_register_id)
        .order_by(Payment.date.desc())
    )
    # ---------- CashTransfers (kassadan kassaga + konvertatsiya) ----------
    tq = (
        db.query(CashTransfer)
        .options(joinedload(CashTransfer.from_cash), joinedload(CashTransfer.to_cash))
        .filter(
            or_(
                CashTransfer.from_cash_id == cash_register_id,
                CashTransfer.to_cash_id == cash_register_id,
            ),
            CashTransfer.status.in_(("in_transit", "completed")),
        )
        .order_by(CashTransfer.date.desc())
    )
    period_opening_balance = None
    period_closing_balance = None
    from app.services.finance_service import cash_balance_formula as _cbf
    if (date_from or "").strip():
        try:
            df = datetime.strptime(str(date_from).strip()[:10], "%Y-%m-%d").date()
            q = q.filter(Payment.date >= df)
            tq = tq.filter(CashTransfer.date >= df)
            period_opening_balance, _, _ = _cbf(db, cash_register_id, as_of_date=df - timedelta(days=1))
        except ValueError:
            pass
    if (date_to or "").strip():
        try:
            dt = datetime.strptime(str(date_to).strip()[:10], "%Y-%m-%d").date()
            cutoff = datetime.combine(dt + timedelta(days=1), datetime.min.time())
            q = q.filter(Payment.date < cutoff)
            tq = tq.filter(CashTransfer.date < cutoff)
            period_closing_balance, _, _ = _cbf(db, cash_register_id, as_of_date=dt)
        except ValueError:
            pass

    # CashTransfer'ni Payment-compatible row ga aylantirish
    from types import SimpleNamespace
    transfers_all = tq.all()
    transfer_rows = []
    for t in transfers_all:
        is_out = t.from_cash_id == cash_register_id
        # Konvertatsiya = cross-currency (to_amount + exchange_rate aniq farqli)
        is_conv = (
            t.exchange_rate is not None
            and t.to_amount is not None
            and abs(float(t.to_amount or 0) - float(t.amount or 0)) > 0.01
        )
        # Bu kassa nuqtai-nazaridan amount: chiqimda from amount, kirimda to_amount (yoki amount agar to_amount=NULL)
        amount = float(t.amount or 0) if is_out else float(t.to_amount or t.amount or 0)
        other = t.to_cash if is_out else t.from_cash
        other_name = other.name if other else "—"
        label = "Konvertatsiya" if is_conv else "Kassadan kassaga"
        arrow = "→" if is_out else "←"
        if is_conv and not is_out:
            # kirimda konvertatsiya: from valyuta + rate ko'rsatish
            rate_str = f" (kurs {t.exchange_rate:g})"
        else:
            rate_str = ""
        desc = f"{label}: {arrow} {other_name}{rate_str} (#{t.number or t.id})"
        transfer_rows.append(SimpleNamespace(
            id=f"T-{t.id}",
            date=t.date,
            type="expense" if is_out else "income",
            amount=amount,
            partner=None,
            description=desc,
            category="transfer",
            payment_type="naqd",
            status=t.status,
            is_transfer=True,
            transfer_id=t.id,
            transfer_number=t.number,
        ))

    # Merge Payment + transfer rows, sort by date desc
    payments_all = q.all()
    all_rows = list(payments_all) + transfer_rows
    all_rows.sort(key=lambda r: (r.date or datetime.min), reverse=True)

    per_page = 100
    page = max(1, int(page)) if page else 1
    total_count = len(all_rows)
    total_pages = max(1, (total_count + per_page - 1) // per_page) if total_count else 1
    page = min(page, total_pages)
    payments = all_rows[(page - 1) * per_page : page * per_page]
    # Har payment uchun bog'langan ExpenseDoc id ni topish (HD-... raqamlari klikab bo'lishi uchun)
    # CashTransfer rowlarni o'tkazib yuborish (ularda ExpenseDoc yo'q)
    payment_ids = [p.id for p in payments if p.id and not getattr(p, 'is_transfer', False)]
    expense_doc_by_payment = {}
    if payment_ids:
        for ed in db.query(ExpenseDoc.id, ExpenseDoc.payment_id, ExpenseDoc.number).filter(
            ExpenseDoc.payment_id.in_(payment_ids)
        ).all():
            expense_doc_by_payment[ed.payment_id] = {"id": ed.id, "number": ed.number}
    filter_date_from = str(date_from or "").strip()[:10] if date_from else ""
    filter_date_to = str(date_to or "").strip()[:10] if date_to else ""
    # Sahifa darajasidagi yig'indi (Payment + CashTransfer)
    total_income = sum(float(p.amount or 0) for p in payments if getattr(p, "type", None) == "income")
    total_expense = sum(float(p.amount or 0) for p in payments if getattr(p, "type", None) == "expense")
    # Davr darajasidagi yig'indi (hamma sahifalar bo'yicha, filter doirasida)
    period_total_income = sum(float(p.amount or 0) for p in all_rows if getattr(p, "type", None) == "income")
    period_total_expense = sum(float(p.amount or 0) for p in all_rows if getattr(p, "type", None) == "expense")
    parts = []
    if filter_date_from:
        parts.append(f"date_from={filter_date_from}")
    if filter_date_to:
        parts.append(f"date_to={filter_date_to}")
    pagination_query = "&".join(parts)
    return templates.TemplateResponse("finance/kassa_detail.html", {
        "request": request,
        "cash": cash,
        "payments": payments,
        "expense_doc_by_payment": expense_doc_by_payment,
        "filter_date_from": filter_date_from,
        "filter_date_to": filter_date_to,
        "total_income": total_income,
        "total_expense": total_expense,
        "total_income_all_time": total_income_all_time,
        "total_expense_all_time": total_expense_all_time,
        "computed_balance": computed_balance,
        "stored_balance": stored_balance,
        "balance_mismatch": balance_mismatch,
        "period_opening_balance": period_opening_balance,
        "period_closing_balance": period_closing_balance,
        "period_total_income": period_total_income,
        "period_total_expense": period_total_expense,
        "page": page,
        "per_page": per_page,
        "total_count": total_count,
        "total_pages": total_pages,
        "pagination_query": pagination_query,
        "current_user": current_user,
        "page_title": (cash.name or "Kassa") + " — kirim/chiqimlar",
    })


@router.post("/payment")
async def finance_payment_post(
    request: Request,
    type: str = Form(...),
    amount: float = Form(...),
    cash_register_id: int = Form(...),
    partner_id: Optional[int] = Form(None),
    description: str = Form(""),
    payment_date: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    ensure_payments_status_column(db)
    if type not in ("income", "expense"):
        return RedirectResponse(url="/finance?error=type", status_code=303)
    cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash:
        return RedirectResponse(url="/finance?error=cash", status_code=303)
    amount = float(amount)
    if amount <= 0:
        return RedirectResponse(url="/finance?error=amount", status_code=303)
    if not partner_id or int(partner_id) <= 0:
        from urllib.parse import quote
        return RedirectResponse(url="/finance?error=" + quote("Kontragent tanlanmagan!"), status_code=303)
    # Dublikat himoyasi: 3 daqiqa ichida bir xil tur+summa+kassa+mijoz
    from datetime import timedelta
    three_min_ago = datetime.now() - timedelta(minutes=3)
    dup_q = db.query(Payment).filter(
        Payment.type == type,
        Payment.amount == amount,
        Payment.cash_register_id == cash_register_id,
        Payment.created_at >= three_min_ago,
    )
    if partner_id and int(partner_id) > 0:
        dup_q = dup_q.filter(Payment.partner_id == int(partner_id))
    if dup_q.first():
        from urllib.parse import quote
        return RedirectResponse(url="/finance?error=" + quote("Oxirgi 3 daqiqada aynan shu to'lov yaratilgan. Takroriy bo'lsa, biroz kuting."), status_code=303)
    pid = None
    if partner_id is not None and int(partner_id) > 0:
        p = db.query(Partner).filter(Partner.id == int(partner_id)).first()
        if p:
            pid = p.id
    # Sana: foydalanuvchi tanlagan sana yoki bugun
    pay_dt = datetime.now()
    if payment_date:
        try:
            d = datetime.strptime(payment_date, "%Y-%m-%d")
            pay_dt = datetime.combine(d.date(), datetime.now().time())
        except ValueError:
            pass
    pay_date_str = pay_dt.strftime('%Y%m%d')
    pay_date_start = pay_dt.replace(hour=0, minute=0, second=0, microsecond=0)
    pay_date_end = pay_dt.replace(hour=23, minute=59, second=59, microsecond=999999)
    today_str = datetime.now().strftime('%Y%m%d')
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    # H5 fix: count()+1 fallback emas, prefiks bo'yicha MAX(suffix)+1 (uniform helper)
    pay_number = next_doc_number(db, Payment, f"PAY-{pay_date_str}-")
    desc = (description or "").strip() or ("Kirim" if type == "income" else "Chiqim")
    payment = Payment(
        number=pay_number,
        date=pay_dt,
        type=type,
        cash_register_id=cash_register_id,
        partner_id=pid,
        order_id=None,
        amount=amount,
        payment_type="cash",
        category="other",
        description=desc,
        user_id=current_user.id if current_user else None,
        status="confirmed",
    )
    db.add(payment)
    db.flush()
    _payment_apply_balance(db, payment, 1)
    log_action(db, user=current_user, action="create", entity_type="payment",
               entity_id=payment.id, entity_number=payment.number,
               details=f"Tur: {type}, Summa: {amount:,.0f}, Partner: {pid or 'yo`q'}",
               ip_address=request.client.host if request.client else "")
    db.commit()
    try:
        from app.bot.services.audit_watchdog import audit_payment
        audit_payment(payment.id)
    except Exception:
        pass
    return RedirectResponse(url="/finance?success=1", status_code=303)


def _payment_apply_balance(db: Session, payment: Payment, sign: int):
    """Kassa va kontragent balanslarini yangilash.

    `sign` — eski inkremental API bilan moslik uchun (endi partner balansiga ta'sir qilmaydi).
    Kontragent balansi recompute bilan hujjatlardan qayta hisoblanadi (USD konvertatsiya bilan #7/#8),
    kassa balansi formula bilan sinxronlanadi. Chaqiruvchi to'lov statusini (confirmed/cancelled)
    DB'da o'rnatib bo'lgach chaqiradi.
    """
    _sync_cash_balance(db, payment.cash_register_id)
    if payment.partner_id:
        from app.services.partner_balance_service import recompute_partner_balance
        db.flush()
        recompute_partner_balance(db, payment.partner_id,
                                  reason=f"payment_{payment.type or 'unknown'}",
                                  ref=payment.number)


@router.post("/payment/{payment_id}/confirm")
async def finance_payment_confirm(
    payment_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    if is_period_closed(db, payment.date):
        return RedirectResponse(url="/finance?error=period_closed", status_code=303)
    # Atomik UPDATE WHERE — double-confirm xavfini oldini olish
    from sqlalchemy import text as _text
    claim = db.execute(
        _text("UPDATE payments SET status='confirmed' WHERE id=:id AND (status IS NULL OR status != 'confirmed')"),
        {"id": payment_id}
    )
    if claim.rowcount == 0:
        return RedirectResponse(url="/finance?msg=already_confirmed", status_code=303)
    db.refresh(payment)  # status atributini yangilash (Python obyekt eskiriksiz qoladi)
    _payment_apply_balance(db, payment, 1)
    db.commit()
    try:
        from app.bot.services.audit_watchdog import audit_payment
        audit_payment(payment.id)
    except Exception:
        pass
    referer = request.headers.get("referer", "/finance")
    sep = "&" if "?" in referer else "?"
    return RedirectResponse(url=f"{referer}{sep}success=confirmed", status_code=303)


@router.post("/payment/{payment_id}/cancel")
async def finance_payment_cancel(
    payment_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    status = getattr(payment, "status", "confirmed")
    if status == "cancelled":
        return RedirectResponse(url="/finance?msg=already_cancelled", status_code=303)
    payment.status = "cancelled"
    _payment_apply_balance(db, payment, -1)
    db.commit()
    referer = request.headers.get("referer", "/finance")
    sep = "&" if "?" in referer else "?"
    return RedirectResponse(url=f"{referer}{sep}success=cancelled", status_code=303)


@router.get("/payment/{payment_id}/edit", response_class=HTMLResponse)
async def finance_payment_edit_page(
    payment_id: int,
    request: Request,
    view: int = 0,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    is_confirmed = getattr(payment, "status", "confirmed") == "confirmed"
    if is_confirmed and not view:
        return RedirectResponse(
            url="/finance?error=" + quote("Tasdiqlangan to'lovni tahrirlash mumkin emas. Avval tasdiqni bekor qiling. Ko'rish uchun ko'z tugmasini bosing."),
            status_code=303,
        )
    partners = db.query(Partner).filter(Partner.is_active == True).order_by(Partner.name).all()
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).all()
    return templates.TemplateResponse("finance/payment_edit.html", {
        "request": request,
        "payment": payment,
        "partners": partners,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "read_only": is_confirmed,
        "page_title": "To'lovni ko'rish" if is_confirmed else "To'lovni tahrirlash",
    })


@router.post("/payment/{payment_id}/edit")
async def finance_payment_edit_post(
    payment_id: int,
    type: str = Form(...),
    amount: float = Form(...),
    cash_register_id: int = Form(...),
    partner_id: Optional[int] = Form(None),
    description: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    if getattr(payment, "status", "confirmed") == "confirmed":
        return RedirectResponse(
            url="/finance?error=" + quote("Tasdiqlangan to'lovni tahrirlash mumkin emas. Avval tasdiqni bekor qiling."),
            status_code=303,
        )
    if type not in ("income", "expense"):
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=type", status_code=303)
    amount = float(amount)
    if amount <= 0:
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=amount", status_code=303)
    cash_new = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash_new:
        return RedirectResponse(url=f"/finance/payment/{payment_id}/edit?error=cash", status_code=303)
    # M7 fix: kassa valyutasi o'zgarsa summa noto'g'ri talqin qilinadi (so'm raqami
    # USD deb yoki aksincha). Avtomatik qayta-konvert noaniq (foydalanuvchi niyati) —
    # xavfsiz yo'l: valyuta o'zgarishini bloklab, o'chirib-qayta yaratishni so'rash.
    old_cr = db.query(CashRegister).filter(CashRegister.id == payment.cash_register_id).first()
    old_curr = (getattr(old_cr, "currency", None) or "UZS").upper() if old_cr else "UZS"
    new_curr = (getattr(cash_new, "currency", None) or "UZS").upper()
    if old_curr != new_curr:
        return RedirectResponse(
            url=f"/finance/payment/{payment_id}/edit?error=" + quote(
                f"Kassa valyutasi farq qiladi ({old_curr}->{new_curr}). Summa noto'g'ri "
                f"talqin qilinmasligi uchun to'lovni o'chirib qayta yarating."
            ),
            status_code=303,
        )
    pid = None
    if partner_id is not None and int(partner_id) > 0:
        p = db.query(Partner).filter(Partner.id == int(partner_id)).first()
        if p:
            pid = p.id
    payment.type = type
    payment.amount = amount
    payment.cash_register_id = cash_register_id
    payment.partner_id = pid
    payment.description = (description or "").strip() or ("Kirim" if type == "income" else "Chiqim")
    # Tahrirdan keyin avtomatik qayta tasdiqlash — balansga qaytarish
    was_cancelled = (getattr(payment, "status", "confirmed") == "cancelled")
    if was_cancelled:
        payment.status = "confirmed"
        _payment_apply_balance(db, payment, +1)
    db.commit()
    return RedirectResponse(url="/finance?success=edited", status_code=303)


@router.post("/payment/{payment_id}/delete")
async def finance_payment_delete(
    payment_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """To'lovni o'chirish (faqat cancelled). Kassa balansi avtomatik qayta hisoblanadi."""
    from app.services.payment_service import delete_payment_atomic
    from app.services.document_service import DocumentError
    payment = db.query(Payment).filter(Payment.id == payment_id).first()
    if not payment:
        raise HTTPException(status_code=404, detail="To'lov topilmadi")
    try:
        delete_payment_atomic(db, payment)
    except DocumentError as e:
        return RedirectResponse(
            url="/finance?error=" + quote(str(e)), status_code=303,
        )
    return RedirectResponse(url="/finance?success=deleted", status_code=303)


@router.get("/expense-types", response_class=HTMLResponse)
async def finance_expense_types_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    types_list = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.name).all()
    if not types_list and db.query(ExpenseType).count() == 0:
        for name, cat in [
            ("ish haqqi", "Ishlab chiqarish xarajatlari"),
            ("ishxona harajati", "Ishlab chiqarish xarajatlari"),
            ("karobka yasatishga", "Ishlab chiqarish xarajatlari"),
            ("oziq ovqatga", "Ma'muriy xarajatlar"),
            ("Yolkiro", "Ma'muriy xarajatlar"),
        ]:
            db.add(ExpenseType(name=name, category=cat, is_active=True))
        db.commit()
        types_list = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.name).all()
    return templates.TemplateResponse("finance/expense_types.html", {
        "request": request,
        "expense_types": types_list,
        "current_user": current_user,
        "page_title": "Harajat turlari",
    })


@router.post("/expense-types/add")
async def finance_expense_type_add(
    name: str = Form(...),
    category: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    name = (name or "").strip()
    if not name:
        return RedirectResponse(url="/finance/expense-types?error=name", status_code=303)
    db.add(ExpenseType(name=name, category=(category or "").strip() or None, is_active=True))
    db.commit()
    return RedirectResponse(url="/finance/expense-types", status_code=303)


@router.post("/expense-types/edit/{etype_id}")
async def finance_expense_type_edit(
    etype_id: int,
    name: str = Form(...),
    category: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    e = db.query(ExpenseType).filter(ExpenseType.id == etype_id).first()
    if not e:
        raise HTTPException(status_code=404, detail="Harajat turi topilmadi")
    e.name = (name or "").strip() or e.name
    e.category = (category or "").strip() or None
    db.commit()
    return RedirectResponse(url="/finance/expense-types", status_code=303)


@router.post("/expense-types/delete/{etype_id}")
async def finance_expense_type_delete(
    etype_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    e = db.query(ExpenseType).filter(ExpenseType.id == etype_id).first()
    if not e:
        raise HTTPException(status_code=404, detail="Harajat turi topilmadi")
    e.is_active = False
    db.commit()
    return RedirectResponse(url="/finance/expense-types", status_code=303)


@router.get("/harajat/hujjat/new", response_class=HTMLResponse)
async def finance_harajat_hujjat_new(
    request: Request,
    force_new: int = 0,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    # Draft cheklov: bitta foydalanuvchi 1 ta ochiq HD draft
    from app.utils.draft_check import redirect_to_draft
    redirect = redirect_to_draft(
        db, ExpenseDoc,
        edit_url_template="/finance/harajat/hujjat/{id}",
        user_role=getattr(current_user, "role", "") or "",
        force_new=bool(force_new),
        message="Sizda ochiq harajat hujjati qoralamasi bor — avval uni tugating yoki bekor qiling.",
        user_id=current_user.id,
    )
    if redirect:
        return redirect
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    directions = db.query(Direction).filter(Direction.is_active == True).order_by(Direction.name).all() if hasattr(Direction, "is_active") else db.query(Direction).order_by(Direction.name).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    expense_types = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.name).all()
    doc_date = datetime.now().date()
    return templates.TemplateResponse("finance/harajat_hujjat_form.html", {
        "request": request,
        "doc": None,
        "doc_date": doc_date,
        "cash_registers": cash_registers,
        "directions": directions,
        "departments": departments,
        "expense_types": expense_types,
        "current_user": current_user,
        "page_title": "Harajat hujjati — yaratish",
    })


@router.get("/harajat/hujjat/{doc_id}", response_class=HTMLResponse)
async def finance_harajat_hujjat_edit(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    doc = db.query(ExpenseDoc).options(
        joinedload(ExpenseDoc.items).joinedload(ExpenseDocItem.expense_type),
        joinedload(ExpenseDoc.cash_register),
        joinedload(ExpenseDoc.direction),
        joinedload(ExpenseDoc.department),
    ).filter(ExpenseDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Harajat hujjati topilmadi")
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    directions = db.query(Direction).filter(Direction.is_active == True).order_by(Direction.name).all() if hasattr(Direction, "is_active") else db.query(Direction).order_by(Direction.name).all()
    departments = db.query(Department).filter(Department.is_active == True).order_by(Department.name).all()
    expense_types = db.query(ExpenseType).filter(ExpenseType.is_active == True).order_by(ExpenseType.name).all()
    if doc.date:
        doc_date = doc.date.date() if hasattr(doc.date, "date") and callable(getattr(doc.date, "date")) else doc.date
    else:
        doc_date = datetime.now().date()
    # Oylik to'lovi item bormi? "Oylik to'lovi YYYY-MM" pattern
    salary_rows = []
    salary_year = None
    salary_month = None
    for item in (doc.items or []):
        m = _re.search(r"Oylik to'lovi\s*(\d{4})-(\d{1,2})", item.description or "")
        if m:
            salary_year = int(m.group(1))
            salary_month = int(m.group(2))
            break
    if salary_year and salary_month:
        sals = db.query(Salary).options(joinedload(Salary.employee)).filter(
            Salary.year == salary_year, Salary.month == salary_month
        ).order_by(Salary.employee_id).all()
        for s in sals:
            if (s.total or 0) <= 0 and (s.paid or 0) <= 0:
                continue
            salary_rows.append({
                "employee": s.employee,
                "base_salary": float(s.base_salary or 0),
                "bonus": float(s.bonus or 0),
                "deduction": float(s.deduction or 0),
                "advance_deduction": float(s.advance_deduction or 0),
                "total": float(s.total or 0),
                "paid": float(s.paid or 0),
                "status": s.status or "pending",
            })
    return templates.TemplateResponse("finance/harajat_hujjat_form.html", {
        "request": request,
        "doc": doc,
        "doc_date": doc_date,
        "cash_registers": cash_registers,
        "directions": directions,
        "departments": departments,
        "expense_types": expense_types,
        "salary_rows": salary_rows,
        "salary_year": salary_year,
        "salary_month": salary_month,
        "current_user": current_user,
        "page_title": f"Harajat hujjati #{doc.number or doc_id}",
    })


def _get_salary_doc_data(db, doc_id):
    """Oylik to'lovi hujjati ma'lumotlari (export uchun)."""
    doc = db.query(ExpenseDoc).options(
        joinedload(ExpenseDoc.items),
        joinedload(ExpenseDoc.cash_register),
    ).filter(ExpenseDoc.id == doc_id).first()
    if not doc:
        return None, None, None, None
    salary_year = None
    salary_month = None
    for item in (doc.items or []):
        m = _re.search(r"Oylik to'lovi\s*(\d{4})-(\d{1,2})", item.description or "")
        if m:
            salary_year = int(m.group(1))
            salary_month = int(m.group(2))
            break
    if not (salary_year and salary_month):
        return doc, None, None, None
    sals = db.query(Salary).options(joinedload(Salary.employee)).filter(
        Salary.year == salary_year, Salary.month == salary_month
    ).order_by(Salary.employee_id).all()
    rows = []
    for s in sals:
        if (s.total or 0) <= 0 and (s.paid or 0) <= 0:
            continue
        rows.append(s)
    return doc, rows, salary_year, salary_month


@router.get("/harajat/hujjat/{doc_id}/salary-export/excel")
async def finance_salary_export_excel(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik to'lovi hujjatini Excel ga eksport qilish."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    doc, rows, y, m = _get_salary_doc_data(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if not rows:
        raise HTTPException(status_code=400, detail="Oylik ro'yxati topilmadi")
    wb = Workbook()
    ws = wb.active
    ws.title = f"Oylik {y}-{m:02d}"
    # Header
    ws['A1'] = f"Oylik to'lovi hujjati — {doc.number}"
    ws['A1'].font = Font(bold=True, size=14)
    ws.merge_cells('A1:I1')
    ws['A2'] = f"{y}-yil {m}-oy | Kassa: {doc.cash_register.name if doc.cash_register else '—'} | Sana: {doc.date.strftime('%d.%m.%Y') if doc.date else '—'}"
    ws.merge_cells('A2:I2')
    # Column headers
    headers = ["#", "Xodim", "Kod", "Lavozim", "Oylik asos", "Bonus", "Ushlab qolish", "Avans", "Jami"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    # Rows
    total_sum = 0
    for i, s in enumerate(rows, 1):
        ws.cell(row=4 + i, column=1, value=i)
        ws.cell(row=4 + i, column=2, value=s.employee.full_name if s.employee else "—")
        ws.cell(row=4 + i, column=3, value=s.employee.code if s.employee else "")
        ws.cell(row=4 + i, column=4, value=s.employee.position if s.employee and s.employee.position else "")
        ws.cell(row=4 + i, column=5, value=float(s.base_salary or 0))
        ws.cell(row=4 + i, column=6, value=float(s.bonus or 0))
        ws.cell(row=4 + i, column=7, value=float(s.deduction or 0))
        ws.cell(row=4 + i, column=8, value=float(s.advance_deduction or 0))
        ws.cell(row=4 + i, column=9, value=float(s.total or 0))
        total_sum += float(s.total or 0)
    # Total row
    tr = 4 + len(rows) + 1
    ws.cell(row=tr, column=1, value="Jami:").font = Font(bold=True)
    ws.merge_cells(start_row=tr, start_column=1, end_row=tr, end_column=8)
    ws.cell(row=tr, column=1).alignment = Alignment(horizontal="right")
    ws.cell(row=tr, column=9, value=total_sum).font = Font(bold=True)
    # Column widths
    widths = [5, 30, 8, 20, 14, 12, 14, 14, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w
    # Number format
    for row in ws.iter_rows(min_row=5, max_row=tr, min_col=5, max_col=9):
        for cell in row:
            cell.number_format = '#,##0'
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"oylik_{y}_{m:02d}_{doc.number}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/harajat/hujjat/{doc_id}/salary-export/word")
async def finance_salary_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Oylik to'lovi hujjatini Word ga eksport qilish."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    doc, rows, y, m = _get_salary_doc_data(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if not rows:
        raise HTTPException(status_code=400, detail="Oylik ro'yxati topilmadi")
    d = Document()
    d.add_heading(f"Oylik to'lovi hujjati — {doc.number}", 0)
    p = d.add_paragraph(f"{y}-yil {m}-oy | Kassa: {doc.cash_register.name if doc.cash_register else '—'} | Sana: {doc.date.strftime('%d.%m.%Y') if doc.date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Table
    table = d.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Xodim", "Lavozim", "Oylik asos", "Bonus", "Ushlab", "Avans", "Jami"]):
        hdr[i].text = h
    total_sum = 0
    for i, s in enumerate(rows, 1):
        r = table.add_row().cells
        r[0].text = str(i)
        r[1].text = (s.employee.full_name or "—") if s.employee else "—"
        r[2].text = (s.employee.position or "—") if s.employee and s.employee.position else "—"
        r[3].text = f"{float(s.base_salary or 0):,.0f}"
        r[4].text = f"{float(s.bonus or 0):,.0f}"
        r[5].text = f"{float(s.deduction or 0):,.0f}"
        r[6].text = f"{float(s.advance_deduction or 0):,.0f}"
        r[7].text = f"{float(s.total or 0):,.0f}"
        total_sum += float(s.total or 0)
    r = table.add_row().cells
    r[0].text = ""; r[1].text = ""; r[2].text = ""; r[3].text = ""
    r[4].text = ""; r[5].text = ""; r[6].text = "Jami:"; r[7].text = f"{total_sum:,.0f}"
    buf = BytesIO()
    d.save(buf)
    buf.seek(0)
    filename = f"oylik_{y}_{m:02d}_{doc.number}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.post("/harajat/hujjat/save")
async def finance_harajat_hujjat_save(
    request: Request,
    doc_id: Optional[int] = Form(None),
    date: str = Form(...),
    cash_register_id: int = Form(...),
    direction_id: Optional[int] = Form(None),
    department_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    cash = db.query(CashRegister).filter(CashRegister.id == cash_register_id).first()
    if not cash:
        return RedirectResponse(url="/finance/harajatlar?error=cash", status_code=303)
    try:
        _d = datetime.strptime(str(date).strip()[:10], "%Y-%m-%d").date()
        # Tanlangan sanaga joriy vaqtni qo'shamiz (faqat sana = 00:00 bo'lib qolmasin, payment bilan mos)
        doc_date = datetime.combine(_d, datetime.now().time())
    except ValueError:
        doc_date = datetime.now()
    form = await request.form()
    ids = form.getlist("expense_type_id")
    amounts = form.getlist("amount")
    descriptions = form.getlist("description")
    if doc_id:
        doc = db.query(ExpenseDoc).filter(ExpenseDoc.id == doc_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Hujjat topilmadi")
        if doc.status == "confirmed":
            return RedirectResponse(url="/finance/harajatlar?error=confirmed", status_code=303)
    else:
        # Draft cheklov: yangi HD yaratish oldidan tekshiruv (defense in depth)
        from app.utils.draft_check import redirect_to_draft
        _redirect = redirect_to_draft(
            db, ExpenseDoc,
            edit_url_template="/finance/harajat/hujjat/{id}",
            user_role=getattr(current_user, "role", "") or "",
            message="Sizda ochiq harajat hujjati qoralamasi bor.",
            user_id=current_user.id,
        )
        if _redirect:
            return _redirect
        doc = ExpenseDoc(
            number=_next_expense_doc_number(db),
            date=doc_date,
            cash_register_id=cash_register_id,
            direction_id=int(direction_id) if direction_id and int(direction_id) > 0 else None,
            department_id=int(department_id) if department_id and int(department_id) > 0 else None,
            status="draft",
            total_amount=0,
            user_id=current_user.id if current_user else None,
        )
        db.add(doc)
        db.flush()
    doc.date = doc_date
    doc.cash_register_id = cash_register_id
    doc.direction_id = int(direction_id) if direction_id and int(direction_id) > 0 else None
    doc.department_id = int(department_id) if department_id and int(department_id) > 0 else None
    doc.user_id = current_user.id if current_user else None
    for it in list(doc.items):
        db.delete(it)
    db.flush()
    total = 0.0
    for i in range(max(len(ids), len(amounts))):
        et_id = int(ids[i]) if i < len(ids) and str(ids[i]).strip().isdigit() else None
        amt = float(amounts[i]) if i < len(amounts) and str(amounts[i]).strip() else 0
        desc = (descriptions[i] if i < len(descriptions) else "").strip() or None
        if et_id and amt > 0:
            et = db.query(ExpenseType).filter(ExpenseType.id == et_id).first()
            if et:
                db.add(ExpenseDocItem(expense_doc_id=doc.id, expense_type_id=et_id, amount=amt, description=desc))
                total += amt
    doc.total_amount = total
    db.commit()
    return RedirectResponse(url=f"/finance/harajat/hujjat/{doc.id}", status_code=303)


@router.post("/harajat/hujjat/{doc_id}/tasdiqlash")
async def finance_harajat_hujjat_tasdiqlash(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    ensure_payments_status_column(db)
    try:
        doc = db.query(ExpenseDoc).options(joinedload(ExpenseDoc.items)).filter(ExpenseDoc.id == doc_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Harajat hujjati topilmadi")
        if doc.status == "confirmed":
            return RedirectResponse(url="/finance/harajatlar?error=already_confirmed", status_code=303)
        if not doc.items:
            return RedirectResponse(url="/finance/harajatlar?error=no_items", status_code=303)
        if not getattr(doc, "cash_register_id", None):
            return RedirectResponse(url="/finance/harajatlar?error=no_cash", status_code=303)
        total = sum(getattr(it, "amount", 0) or 0 for it in doc.items)
        if total <= 0:
            return RedirectResponse(url="/finance/harajatlar?error=no_amount", status_code=303)
        claim = db.execute(
            text("UPDATE expense_docs SET status = 'confirmed', total_amount = :tot WHERE id = :id AND status != 'confirmed'"),
            {"tot": total, "id": doc_id}
        )
        if claim.rowcount == 0:
            db.rollback()
            return RedirectResponse(url="/finance/harajatlar?error=already_confirmed", status_code=303)
        pay_number = f"PAY-{datetime.now().strftime('%Y%m%d%H%M%S')}-D{doc_id}"
        payment_date = datetime.now()
        if getattr(doc, "date", None):
            d = doc.date
            if hasattr(d, "date") and callable(getattr(d, "date")):
                payment_date = datetime.combine(d.date(), datetime.now().time())
            else:
                payment_date = d
        uid = getattr(current_user, "id", None) if current_user else None
        if uid is None:
            first_user = db.query(User).order_by(User.id).first()
            uid = first_user.id if first_user else None
        if uid is None:
            db.rollback()
            return RedirectResponse(url="/finance/harajatlar?error=no_user", status_code=303)
        payment = Payment(
            number=pay_number,
            date=payment_date,
            type="expense",
            cash_register_id=doc.cash_register_id,
            partner_id=None,
            order_id=None,
            amount=total,
            payment_type="cash",
            category="expense_doc",
            description=f"Harajat hujjati #{doc.number or doc_id}",
            user_id=uid,
            status="confirmed",
        )
        db.add(payment)
        db.flush()
        db.execute(
            text("UPDATE expense_docs SET payment_id = :pid WHERE id = :id"),
            {"pid": payment.id, "id": doc_id}
        )
        _sync_cash_balance(db, doc.cash_register_id)
        db.commit()
        # Telegram bildirish (ELYA CLASSIC — real-time)
        try:
            from app.bot.services.notifier import notify_expense
            notify_expense(doc.number or f"#{doc_id}", total, "")
        except Exception:
            pass
        try:
            from app.bot.services.audit_watchdog import audit_expense, audit_payment
            audit_expense(doc_id)
            audit_payment(payment.id)
        except Exception:
            pass
        return RedirectResponse(url="/finance/harajatlar?success=confirmed", status_code=303)
    except HTTPException:
        raise
    except IntegrityError:
        db.rollback()
        return RedirectResponse(url="/finance/harajatlar?error=duplicate", status_code=303)
    except Exception:
        db.rollback()
        return RedirectResponse(url="/finance/harajatlar?error=save_error", status_code=303)


@router.post("/harajat/hujjat/{doc_id}/revert")
async def finance_harajat_hujjat_revert(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Harajat hujjati tasdiqini bekor qilish (faqat admin) — payment bekor, kassa qaytariladi"""
    doc = db.query(ExpenseDoc).filter(ExpenseDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status != "confirmed":
        return RedirectResponse(url="/finance/harajatlar?error=not_confirmed", status_code=303)
    # Payment ni bekor qilish
    if doc.payment_id:
        payment = db.query(Payment).filter(Payment.id == doc.payment_id).first()
        if payment:
            payment.status = "cancelled"
    db.execute(
        text("UPDATE expense_docs SET status = 'draft', payment_id = NULL WHERE id = :id"),
        {"id": doc_id}
    )
    if doc.cash_register_id:
        _sync_cash_balance(db, doc.cash_register_id)
    db.commit()
    return RedirectResponse(url="/finance/harajatlar?success=reverted", status_code=303)


@router.post("/harajat/hujjat/{doc_id}/delete")
async def finance_harajat_hujjat_delete(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Harajat hujjatini o'chirish (faqat qoralama, faqat admin)"""
    doc = db.query(ExpenseDoc).filter(ExpenseDoc.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if doc.status == "confirmed":
        return RedirectResponse(url="/finance/harajatlar?error=confirmed_cant_delete", status_code=303)
    doc.status = "deleted"
    db.commit()
    return RedirectResponse(url="/finance/harajatlar?success=deleted", status_code=303)


# ==========================================
# KASSADAN KASSAGA O'TKAZISH (/cash/transfers)
# ==========================================

@cash_router.get("/transfiers")
async def cash_transfiers_redirect():
    """Yozuv xatosi: transfiers -> transfers (ro'yxatga yo'naltirish)."""
    return RedirectResponse(url="/cash/transfers", status_code=301)


def _user_owns_cash_register(user, cash_register_id):
    """Foydalanuvchi shu kassaga tegishlimi?"""
    if not user or not cash_register_id:
        return False
    if getattr(user, "cash_register_id", None) == cash_register_id:
        return True
    for cr in (getattr(user, "cash_registers_list", None) or []):
        if cr.id == cash_register_id:
            return True
    return False


@cash_router.get("/transfers", response_class=HTMLResponse)
async def cash_transfers_list(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inkasatsiya hujjatlari ro'yxati"""
    role = (getattr(current_user, "role", None) or "").strip().lower()
    q = (
        db.query(CashTransfer)
        .options(
            joinedload(CashTransfer.from_cash),
            joinedload(CashTransfer.to_cash),
            joinedload(CashTransfer.user),
            joinedload(CashTransfer.approved_by),
        )
        .order_by(CashTransfer.created_at.desc())
    )
    # Sotuvchi faqat o'z kassasiga tegishli hujjatlarni ko'radi
    if role not in ("admin", "manager", "menejer", "rahbar", "raxbar"):
        user_cash_ids = []
        if getattr(current_user, "cash_register_id", None):
            user_cash_ids.append(current_user.cash_register_id)
        for cr in (getattr(current_user, "cash_registers_list", None) or []):
            user_cash_ids.append(cr.id)
        if user_cash_ids:
            q = q.filter(CashTransfer.from_cash_id.in_(user_cash_ids))
        else:
            q = q.filter(CashTransfer.id == -1)  # hech narsa ko'rsatma
    transfers = q.limit(100).all()
    return templates.TemplateResponse("cash/transfers_list.html", {
        "request": request,
        "transfers": transfers,
        "current_user": current_user,
        "page_title": "Inkasatsiya — kassadan kassaga o'tkazish",
    })


@cash_router.get("/transfers/new", response_class=HTMLResponse)
async def cash_transfer_new(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Yangi inkasatsiya hujjati yaratish (faqat admin)"""
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    return templates.TemplateResponse("cash/transfer_form.html", {
        "request": request,
        "transfer": None,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "today_str": datetime.now().strftime("%Y-%m-%d"),
        "page_title": "Inkasatsiya yaratish",
    })


@cash_router.get("/balance-on-date")
async def cash_balance_on_date(
    request: Request,
    date: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """AJAX — tanlangan sana oxirigacha barcha kassalar balansini qaytarish."""
    try:
        as_of = datetime.strptime(date[:10], "%Y-%m-%d")
    except (ValueError, TypeError):
        return JSONResponse({"error": "Sana noto'g'ri"}, status_code=400)
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    result = []
    for c in cash_registers:
        bal, _, _ = _cash_balance_formula_at(db, c.id, as_of)
        result.append({"id": c.id, "name": c.name, "balance": bal})
    return JSONResponse({"date": date[:10], "cash_registers": result})


def _cash_balance_formula_at(db: Session, cash_id: int, as_of_date) -> tuple:
    from app.services.finance_service import cash_balance_formula
    return cash_balance_formula(db, cash_id, as_of_date=as_of_date)


@cash_router.post("/transfers/create")
async def cash_transfer_create(
    request: Request,
    from_cash_id: int = Form(...),
    to_cash_id: int = Form(...),
    amount: float = Form(...),
    exchange_rate: Optional[float] = Form(None),
    to_amount: Optional[float] = Form(None),
    note: str = Form(""),
    date: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Admin inkasatsiya hujjat yaratadi — status: pending.

    date — retroaktiv kiritish uchun (YYYY-MM-DD). Agar bo'sh bo'lsa, bugungi sana.

    Cross-currency: agar from va to kassalar har xil valyutada bo'lsa, exchange_rate
    va to_amount majburiy (yoki ExchangeRate jadvalidan avtomatik olinadi).
    """
    if from_cash_id == to_cash_id:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Qayerdan va qayerga kassa bir xil bolmasin."), status_code=303)
    if amount <= 0:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Summa 0 dan katta bolishi kerak."), status_code=303)

    # Cross-currency aniqlash
    from_cash = db.query(CashRegister).filter(CashRegister.id == from_cash_id).first()
    to_cash = db.query(CashRegister).filter(CashRegister.id == to_cash_id).first()
    if not from_cash or not to_cash:
        return RedirectResponse(url="/cash/transfers/new?error=" + quote("Kassa topilmadi."), status_code=303)
    from_curr = (from_cash.currency or "UZS").upper()
    to_curr = (to_cash.currency or "UZS").upper()
    final_rate = None
    final_to_amount = None
    if from_curr != to_curr:
        # Cross-currency: kurs kerak
        rate = exchange_rate
        if not rate or rate <= 0:
            # Avtomatik kurs olish
            from app.services.currency_service import get_rate
            rate = get_rate(db, from_curr, to_curr)
            if rate <= 0:
                return RedirectResponse(url="/cash/transfers/new?error=" + quote(
                    f"Kurs topilmadi ({from_curr}->{to_curr}). Avval /admin/exchange-rates da kurs kiriting."
                ), status_code=303)
        final_rate = float(rate)
        final_to_amount = float(to_amount) if to_amount and to_amount > 0 else round(amount * final_rate, 2)

    now = datetime.now()
    op_date = now
    if date:
        try:
            parsed = datetime.strptime(date[:10], "%Y-%m-%d")
            today = now.date()
            if parsed.date() > today:
                return RedirectResponse(url="/cash/transfers/new?error=" + quote("Kelajakdagi sana bolishi mumkin emas."), status_code=303)
            if is_period_closed(db, parsed):
                return RedirectResponse(url="/cash/transfers/new?error=" + quote(f"{parsed.strftime('%Y-%m')} oyi yopilgan — retroaktiv kiritish mumkin emas."), status_code=303)
            if parsed.date() == today:
                op_date = now
            else:
                op_date = parsed.replace(hour=23, minute=59, second=59)
        except ValueError:
            return RedirectResponse(url="/cash/transfers/new?error=" + quote("Sana formati notogri."), status_code=303)

    # H5 fix: global CashTransfer id+1 emas, prefiks bo'yicha MAX(suffix)+1
    num = next_doc_number(db, CashTransfer, f"KK-{op_date.strftime('%Y%m%d')}-")
    t = CashTransfer(
        number=num,
        date=op_date,
        from_cash_id=from_cash_id,
        to_cash_id=to_cash_id,
        amount=amount,
        exchange_rate=final_rate,
        to_amount=final_to_amount,
        status="pending",
        user_id=current_user.id if current_user else None,
        note=note or None,
    )
    db.add(t)
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{t.id}", status_code=303)


@cash_router.post("/transfers/{transfer_id}/edit")
async def cash_transfer_edit(
    transfer_id: int,
    request: Request,
    from_cash_id: int = Form(...),
    to_cash_id: int = Form(...),
    amount: float = Form(...),
    note: str = Form(""),
    date: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Admin pending statusdagi inkasatsiya hujjatini tahrirlaydi.

    in_transit/completed statusda — avval revert qilish kerak.
    """
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "pending":
        return RedirectResponse(
            url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat 'pending' statusdagi hujjatni tahrirlash mumkin. Avval 'Bekor qilish' bosib pending ga qaytaring."),
            status_code=303,
        )
    if from_cash_id == to_cash_id:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Qayerdan va qayerga kassa bir xil bolmasin."), status_code=303)
    if amount <= 0:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Summa 0 dan katta bolishi kerak."), status_code=303)

    new_date = t.date
    if date:
        try:
            parsed = datetime.strptime(date[:10], "%Y-%m-%d")
            now = datetime.now()
            today = now.date()
            if parsed.date() > today:
                return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Kelajakdagi sana bolishi mumkin emas."), status_code=303)
            if is_period_closed(db, parsed):
                return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote(f"{parsed.strftime('%Y-%m')} oyi yopilgan."), status_code=303)
            # Bugun = hozirgi vaqt; eski kun = 23:59:59. Sana o'zgarmagan bo'lsa, eski vaqt saqlanadi.
            if t.date and parsed.date() == t.date.date():
                new_date = t.date  # Sana o'zgarmadi — vaqtni saqlash
            elif parsed.date() == today:
                new_date = now
            else:
                new_date = parsed.replace(hour=23, minute=59, second=59)
        except ValueError:
            return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Sana formati notogri."), status_code=303)

    # M8 fix: cross-currency o'tkazmada to_amount/exchange_rate'ni yangi from/to/amount
    # bo'yicha QAYTA hisoblash (avval eski qiymatlar qolib, qabul qiluvchi kassa balansi
    # buzilardi). Bir xil valyutaga o'zgartirilsa maydonlar tozalanadi.
    try:
        new_rate, new_to_amount = _transfer_conversion(db, from_cash_id, to_cash_id, amount)
    except ValueError as ex:
        return RedirectResponse(
            url=f"/cash/transfers/{transfer_id}?error=" + quote(
                f"Kurs topilmadi ({ex}). Avval /admin/exchange-rates da kurs kiriting."
            ),
            status_code=303,
        )
    t.from_cash_id = from_cash_id
    t.to_cash_id = to_cash_id
    t.amount = amount
    t.exchange_rate = new_rate
    t.to_amount = new_to_amount
    t.date = new_date
    t.note = note or None
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?edited=1", status_code=303)


@cash_router.post("/transfers/sotuvchi-send")
async def cash_transfer_sotuvchi_send(
    to_cash_id: int = Form(...),
    amount: float = Form(...),
    from_cash_id: Optional[int] = Form(None),
    inkasator_id: Optional[int] = Form(None),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Sotuvchi o'z kassasidan pul yuboradi — darhol in_transit status.
    Admin/manager har qanday kassadan yuborishi mumkin (from_cash_id Form bilan).
    Sotuvchi faqat o'z kassasidan — from_cash_id avtomatik aniqlanadi."""
    role = (getattr(current_user, "role", None) or "").strip().lower()
    is_admin = role in ("admin", "manager", "menejer", "rahbar", "raxbar")

    # From cash aniqlash
    if is_admin and from_cash_id:
        src_id = from_cash_id
    else:
        # Sotuvchi uchun — user's cash
        user_cash_ids = []
        if getattr(current_user, "cash_register_id", None):
            user_cash_ids.append(current_user.cash_register_id)
        for cr in (getattr(current_user, "cash_registers_list", None) or []):
            user_cash_ids.append(cr.id)
        if not user_cash_ids:
            return JSONResponse({"ok": False, "error": "Sizga kassa biriktirilmagan"}, status_code=400)
        if from_cash_id and from_cash_id in user_cash_ids:
            src_id = from_cash_id
        else:
            src_id = user_cash_ids[0]

    if src_id == to_cash_id:
        return JSONResponse({"ok": False, "error": "Manba va qabul kassasi bir xil bolmasin"}, status_code=400)
    if not amount or amount <= 0:
        return JSONResponse({"ok": False, "error": "Summa 0 dan katta bolishi kerak"}, status_code=400)

    from_cash = db.query(CashRegister).filter(CashRegister.id == src_id).with_for_update().first()
    to_cash = db.query(CashRegister).filter(CashRegister.id == to_cash_id).first()
    if not from_cash or not to_cash:
        return JSONResponse({"ok": False, "error": "Kassa topilmadi"}, status_code=404)

    # Balans tekshirish
    if (from_cash.balance or 0) < amount:
        return JSONResponse({
            "ok": False,
            "error": f"Kassada yetarli mablag' yo'q: bor {(from_cash.balance or 0):,.0f}, kerak {amount:,.0f}",
        }, status_code=400)

    # Inkasator izohi
    full_note = note or ""
    if inkasator_id:
        inkasator = db.query(Partner).filter(Partner.id == inkasator_id).first()
        if inkasator:
            ink_note = f"Inkasator: {inkasator.name}"
            full_note = (full_note + " | " + ink_note) if full_note else ink_note

    # H5 fix: global CashTransfer id+1 emas, prefiks bo'yicha MAX(suffix)+1
    num = next_doc_number(db, CashTransfer, f"KK-{datetime.now().strftime('%Y%m%d')}-")
    t = CashTransfer(
        number=num,
        from_cash_id=src_id,
        to_cash_id=to_cash_id,
        amount=float(amount),
        status="in_transit",  # darhol yo'lda
        user_id=current_user.id if current_user else None,
        sent_by_user_id=current_user.id if current_user else None,
        sent_at=datetime.now(),
        note=full_note or None,
    )
    db.add(t)
    db.flush()
    _sync_cash_balance(db, src_id)
    db.commit()
    try:
        from app.bot.services.audit_watchdog import audit_cash_transfer
        audit_cash_transfer(t.id)
    except Exception:
        pass
    return JSONResponse({"ok": True, "id": t.id, "number": t.number})


_INKASATOR_RE = _re.compile(r"Inkasator:\s*([^|]+?)(?:\s*\||$)")


def _extract_inkasator(note: Optional[str]) -> str:
    if not note:
        return ""
    m = _INKASATOR_RE.search(note)
    return (m.group(1).strip() if m else "")


@cash_router.get("/transfers/my-pending")
async def cash_transfers_my_pending(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Sotuvchining kassasiga tegishli inkasatsiya hujjatlari (JSON).

    Qaytaradi:
    - Faol (pending/in_transit) — aksion uchun
    - Oxirgi 7 kun ichidagi completed — tarix (kim tasdiqladi, qachon, inkasator kim)
    Har yozuvga: inkasator nomi (note dan), sent_at, approved_at, foydalanuvchilar."""
    from datetime import timedelta as _td
    from sqlalchemy import and_, or_

    role = (getattr(current_user, "role", None) or "").strip().lower()
    is_admin = role in ("admin", "manager", "menejer", "rahbar", "raxbar")
    recent_cutoff = datetime.now() - _td(days=7)

    if is_admin:
        active_filter = CashTransfer.status.in_(("pending", "in_transit"))
        recent_completed = and_(
            CashTransfer.status == "completed",
            CashTransfer.approved_at >= recent_cutoff,
        )
        transfers = (
            db.query(CashTransfer)
            .options(joinedload(CashTransfer.from_cash), joinedload(CashTransfer.to_cash))
            .filter(or_(active_filter, recent_completed))
            .order_by(CashTransfer.id.desc())
            .limit(30)
            .all()
        )

        def _role_for(t):
            if t.status == "completed":
                return "done"
            return "sender" if t.status == "pending" else "receiver"
    else:
        user_cash_ids = []
        if getattr(current_user, "cash_register_id", None):
            user_cash_ids.append(current_user.cash_register_id)
        for cr in (getattr(current_user, "cash_registers_list", None) or []):
            user_cash_ids.append(cr.id)
        if not user_cash_ids:
            return JSONResponse([])
        active_filter = or_(
            and_(CashTransfer.from_cash_id.in_(user_cash_ids), CashTransfer.status == "pending"),
            and_(CashTransfer.to_cash_id.in_(user_cash_ids), CashTransfer.status.in_(("pending", "in_transit"))),
        )
        recent_completed = and_(
            CashTransfer.status == "completed",
            CashTransfer.approved_at >= recent_cutoff,
            or_(
                CashTransfer.from_cash_id.in_(user_cash_ids),
                CashTransfer.to_cash_id.in_(user_cash_ids),
            ),
        )
        transfers = (
            db.query(CashTransfer)
            .options(joinedload(CashTransfer.from_cash), joinedload(CashTransfer.to_cash))
            .filter(or_(active_filter, recent_completed))
            .order_by(CashTransfer.id.desc())
            .limit(30)
            .all()
        )

        def _role_for(t):
            if t.status == "completed":
                return "done"
            if t.from_cash_id in user_cash_ids and t.status == "pending":
                return "sender"
            return "receiver"

    # User nomlarini oldindan yuklash
    user_ids = set()
    for t in transfers:
        if t.sent_by_user_id:
            user_ids.add(t.sent_by_user_id)
        if t.approved_by_user_id:
            user_ids.add(t.approved_by_user_id)
    users_map = {}
    if user_ids:
        users_map = {
            u.id: (u.username or f"#{u.id}")
            for u in db.query(User).filter(User.id.in_(user_ids)).all()
        }
    result = []
    for t in transfers:
        result.append({
            "id": t.id,
            "number": t.number,
            "amount": t.amount or 0,
            "from_cash": t.from_cash.name if t.from_cash else "?",
            "to_cash": t.to_cash.name if t.to_cash else "?",
            "note": t.note or "",
            "inkasator": _extract_inkasator(t.note),
            "date": t.created_at.strftime("%d.%m.%Y %H:%M") if t.created_at else "",
            "sent_at": t.sent_at.strftime("%d.%m %H:%M") if t.sent_at else "",
            "sent_by": users_map.get(t.sent_by_user_id, "") if t.sent_by_user_id else "",
            "approved_at": t.approved_at.strftime("%d.%m %H:%M") if t.approved_at else "",
            "approved_by": users_map.get(t.approved_by_user_id, "") if t.approved_by_user_id else "",
            "status": t.status,
            "role": _role_for(t),
        })
    return JSONResponse(result)


@cash_router.get("/transfers/{transfer_id}", response_class=HTMLResponse)
async def cash_transfer_view(
    request: Request,
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    transfer = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not transfer:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    cash_registers = db.query(CashRegister).filter(CashRegister.is_active == True).order_by(CashRegister.name).all()
    return templates.TemplateResponse("cash/transfer_form.html", {
        "request": request,
        "transfer": transfer,
        "cash_registers": cash_registers,
        "current_user": current_user,
        "today_str": datetime.now().strftime("%Y-%m-%d"),
        "page_title": f"Inkasatsiya {transfer.number}",
    })


@cash_router.post("/transfers/{transfer_id}/sotuvchi-confirm")
async def cash_transfer_sotuvchi_confirm(
    transfer_id: int,
    request: Request = None,
    inkasator_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Sotuvchi tasdiqlaydi — pulni inkasatorga berdi. Status: pending -> in_transit"""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "pending":
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat kutilayotgan hujjatni tasdiqlash mumkin."), status_code=303)
    # Sotuvchi faqat o'z kassasidan tasdiqlashi mumkin
    if not _user_owns_cash_register(current_user, t.from_cash_id):
        role = (getattr(current_user, "role", None) or "").strip().lower()
        if role not in ("admin", "manager", "menejer"):
            return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Bu kassa sizga tegishli emas."), status_code=303)
    from_cash = db.query(CashRegister).filter(CashRegister.id == t.from_cash_id).first()
    if not from_cash:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Kassa topilmadi."), status_code=303)
    amount = t.amount or 0
    if (from_cash.balance or 0) < amount:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Kassada yetarli mablag yoq."), status_code=303)
    # Atomik UPDATE WHERE — double-confirm xavfini oldini olish
    from sqlalchemy import text as _text
    claim = db.execute(
        _text("UPDATE cash_transfers SET status='in_transit', sent_by_user_id=:uid, sent_at=:at "
              "WHERE id=:id AND status='pending'"),
        {"id": transfer_id, "uid": current_user.id, "at": datetime.now()}
    )
    if claim.rowcount == 0:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Hujjat allaqachon yuborilgan."), status_code=303)
    # Inkasator ma'lumotini saqlash (note ga qo'shish) — atomik UPDATE muvaffaqiyatli bo'lgandan keyin
    if inkasator_id:
        inkasator = db.query(Partner).filter(Partner.id == inkasator_id).first()
        if inkasator:
            db.refresh(t)
            ink_note = f"Inkasator: {inkasator.name}"
            t.note = (t.note + " | " + ink_note) if t.note else ink_note
    db.flush()
    _sync_cash_balance(db, t.from_cash_id)
    db.commit()
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?sent=1", status_code=303)


@cash_router.get("/transfers/{transfer_id}/receipt", response_class=HTMLResponse)
async def cash_transfer_receipt(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Inkasatsiya cheki — chop etish uchun"""
    t = db.query(CashTransfer).options(
        joinedload(CashTransfer.from_cash),
        joinedload(CashTransfer.to_cash),
        joinedload(CashTransfer.user),
        joinedload(CashTransfer.sent_by),
    ).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    html = f"""<!DOCTYPE html>
<html lang="uz">
<head>
<meta charset="UTF-8">
<title>Inkasatsiya cheki {t.number}</title>
<style>
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: 'Segoe UI', Arial, sans-serif; width: 350px; margin: 0 auto; padding: 20px 15px; color: #1a1d21; }}
.header {{ text-align: center; padding-bottom: 15px; border-bottom: 2px dashed #ccc; margin-bottom: 15px; }}
.header h1 {{ font-size: 18px; font-weight: 800; color: #0d6b4b; margin-bottom: 4px; }}
.header .subtitle {{ font-size: 12px; color: #888; }}
.row {{ display: flex; justify-content: space-between; padding: 6px 0; font-size: 13px; border-bottom: 1px solid #f0f0f0; }}
.row .label {{ color: #666; }}
.row .value {{ font-weight: 600; text-align: right; }}
.amount-row {{ padding: 12px 0; margin: 10px 0; border-top: 2px solid #0d6b4b; border-bottom: 2px solid #0d6b4b; }}
.amount-row .label {{ font-size: 14px; font-weight: 700; }}
.amount-row .value {{ font-size: 20px; font-weight: 800; color: #0d6b4b; }}
.footer {{ text-align: center; margin-top: 20px; padding-top: 15px; border-top: 2px dashed #ccc; }}
.footer .sign {{ margin-top: 30px; display: flex; justify-content: space-between; }}
.footer .sign div {{ text-align: center; flex: 1; }}
.footer .sign .line {{ border-bottom: 1px solid #333; height: 25px; margin-bottom: 4px; }}
.footer .sign .name {{ font-size: 10px; color: #888; }}
.status {{ display: inline-block; background: #fff3cd; color: #856404; padding: 3px 10px; border-radius: 6px; font-size: 11px; font-weight: 700; }}
@media print {{ .no-print {{ display: none !important; }} body {{ padding: 10px; }} }}
</style>
</head>
<body>
<div class="no-print" style="margin-bottom:15px;">
    <button onclick="window.print();" style="padding:8px 20px;background:#0d6b4b;color:#fff;border:none;border-radius:8px;font-weight:600;cursor:pointer;">Chop etish</button>
</div>
<div class="header">
    <h1>INKASATSIYA CHEKI</h1>
    <div class="subtitle">Kassadan kassaga pul o'tkazish</div>
</div>
<div class="row"><span class="label">Hujjat №</span><span class="value">{t.number}</span></div>
<div class="row"><span class="label">Sana</span><span class="value">{t.created_at.strftime('%d.%m.%Y %H:%M') if t.created_at else '-'}</span></div>
<div class="row"><span class="label">Qayerdan</span><span class="value">{t.from_cash.name if t.from_cash else '-'}</span></div>
<div class="row"><span class="label">Qayerga</span><span class="value">{t.to_cash.name if t.to_cash else '-'}</span></div>
<div class="row"><span class="label">Status</span><span class="value"><span class="status">{'Yolda' if t.status == 'in_transit' else t.status}</span></span></div>
<div class="row"><span class="label">Jo'natuvchi</span><span class="value">{t.sent_by.full_name or t.sent_by.username if t.sent_by else '-'}</span></div>
<div class="row"><span class="label">Jo'natish vaqti</span><span class="value">{t.sent_at.strftime('%d.%m.%Y %H:%M') if t.sent_at else '-'}</span></div>
{('<div class="row"><span class="label">Izoh</span><span class="value">' + (t.note or '') + '</span></div>') if t.note else ''}
<div class="row amount-row"><span class="label">SUMMA</span><span class="value">{t.amount:,.0f} som</span></div>
<div class="footer">
    <div class="sign">
        <div><div class="line"></div><div class="name">Berdi (sotuvchi)</div></div>
        <div><div class="line"></div><div class="name">Oldi (inkasator)</div></div>
    </div>
</div>
<script>window.onload=function(){{ try {{ window.print(); }} catch(e){{}} }};</script>
</body>
</html>"""
    return HTMLResponse(content=html)


@cash_router.post("/transfers/{transfer_id}/admin-confirm")
async def cash_transfer_admin_confirm(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Qabul qiluvchi (admin yoki to_cash sotuvchisi) tasdiqlaydi.
    Status: in_transit -> completed."""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status != "in_transit":
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Faqat yoldagi hujjatni qabul qilish mumkin."), status_code=303)
    # Admin/manager yoki to_cash ega sotuvchi qabul qila oladi
    role = (getattr(current_user, "role", None) or "").strip().lower()
    is_admin = role in ("admin", "manager", "menejer", "rahbar", "raxbar")
    if not is_admin and not _user_owns_cash_register(current_user, t.to_cash_id):
        return RedirectResponse(
            url=f"/cash/transfers/{transfer_id}?error=" + quote("Bu qabul qiluvchi kassa sizga tegishli emas."),
            status_code=303,
        )
    to_cash = db.query(CashRegister).filter(CashRegister.id == t.to_cash_id).first()
    if not to_cash:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Qabul kassasi topilmadi."), status_code=303)
    # Atomik UPDATE WHERE — double-confirm xavfini oldini olish
    from sqlalchemy import text as _text
    claim = db.execute(
        _text("UPDATE cash_transfers SET status='completed', approved_by_user_id=:uid, approved_at=:at "
              "WHERE id=:id AND status='in_transit'"),
        {"id": transfer_id, "uid": current_user.id, "at": datetime.now()}
    )
    if claim.rowcount == 0:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote("Hujjat allaqachon tasdiqlangan."), status_code=303)
    db.flush()
    _sync_cash_balance(db, t.to_cash_id)
    db.commit()
    try:
        from app.bot.services.audit_watchdog import audit_cash_transfer
        audit_cash_transfer(t.id)
    except Exception:
        pass
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?confirmed=1", status_code=303)


# Alias: /send -> /sotuvchi-confirm (template backward compat)
@cash_router.post("/transfers/{transfer_id}/send")
async def cash_transfer_send_alias(
    transfer_id: int,
    request: Request = None,
    inkasator_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    return await cash_transfer_sotuvchi_confirm(
        transfer_id, request=request, inkasator_id=inkasator_id, db=db, current_user=current_user,
    )


# Eski endpoint nomini saqlaymiz (backward compat)
@cash_router.post("/transfers/{transfer_id}/confirm")
async def cash_transfer_confirm_legacy(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Legacy confirm — statusga qarab sotuvchi yoki admin confirm ga yo'naltiradi"""
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    if t.status == "pending":
        return await cash_transfer_sotuvchi_confirm(transfer_id, db=db, current_user=current_user)
    elif t.status == "in_transit":
        return await cash_transfer_admin_confirm(transfer_id, db=db, current_user=current_user)
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}", status_code=303)


@cash_router.post("/transfers/{transfer_id}/revert")
async def cash_transfer_revert(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    """Tasdiqni bekor qilish (faqat admin)"""
    from app.services.finance_service import revert_cash_transfer_atomic
    from app.services.document_service import DocumentError
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        return RedirectResponse(url="/cash/transfers?error=" + quote("Hujjat topilmadi."), status_code=303)
    try:
        revert_cash_transfer_atomic(db, t)
    except DocumentError as e:
        return RedirectResponse(url=f"/cash/transfers/{transfer_id}?error=" + quote(str(e)), status_code=303)
    return RedirectResponse(url=f"/cash/transfers/{transfer_id}?reverted=1", status_code=303)


@cash_router.post("/transfers/{transfer_id}/delete")
async def cash_transfer_delete(
    transfer_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    from app.services.finance_service import delete_cash_transfer_atomic
    from app.services.document_service import DocumentError
    t = db.query(CashTransfer).filter(CashTransfer.id == transfer_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    try:
        delete_cash_transfer_atomic(db, t)
    except DocumentError as e:
        return RedirectResponse(url="/cash/transfers?error=" + quote(str(e)), status_code=303)
    return RedirectResponse(url="/cash/transfers?deleted=1", status_code=303)
