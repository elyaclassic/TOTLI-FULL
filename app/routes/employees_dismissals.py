"""
Xodimlar — ishdan bo'shatish hujjatlari (Tier C1 1-bosqich).

Manba: employees.py:216-379 (155 qator) dan ajratib olindi.
Endpoint path'lar o'zgarishsiz — URL'lar eskidek ishlashda davom etadi.
"""
import io
from datetime import datetime
from urllib.parse import quote

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fastapi import APIRouter, Request, Depends, Form, HTTPException, Query
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from sqlalchemy.orm import Session, joinedload

from app.core import templates
from app.models.database import get_db, User, Employee, DismissalDoc
from app.deps import require_auth

router = APIRouter(prefix="/employees", tags=["employees-dismissals"])


# --- ISHDAN BO'SHATISH SABABI ---
DISMISSAL_REASONS = [
    ("own_will", "O'z ixtiyori bilan"),
    ("contract_end", "Shartnoma muddati tugadi"),
    ("discipline", "Mehnat intizomini buzgani"),
    ("reduction", "Loyihadan (shtatdan) qisqartirish"),
    ("agreement", "O'zaro kelishuv"),
    ("other", "Boshqa"),
]


@router.get("/dismissal/create", response_class=HTMLResponse)
async def dismissal_create_page(
    request: Request,
    employee_id: int = Query(..., description="Xodim ID"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati yaratish — forma."""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    default_date = datetime.now().date().strftime("%Y-%m-%d")
    return templates.TemplateResponse("employees/dismissal_form.html", {
        "request": request,
        "employee": emp,
        "reasons": DISMISSAL_REASONS,
        "default_date": default_date,
        "current_user": current_user,
        "page_title": "Ishdan bo'shatish",
    })


@router.post("/dismissal/create", response_class=RedirectResponse)
async def dismissal_create_submit(
    request: Request,
    employee_id: int = Form(...),
    doc_date: str = Form(...),
    reason: str = Form(""),
    note: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini yaratadi, xodimni faol emas qiladi."""
    emp = db.query(Employee).filter(Employee.id == employee_id).first()
    if not emp:
        return RedirectResponse(url="/employees?error=Xodim topilmadi", status_code=303)
    if not emp.is_active:
        return RedirectResponse(url="/employees?error=Xodim allaqachon ishdan bo'shatilgan", status_code=303)
    try:
        doc_d = datetime.strptime(doc_date.strip()[:10], "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return RedirectResponse(
            url=f"/dismissal/create?employee_id={employee_id}&error=Noto%27g%27ri sana",
            status_code=303,
        )
    reason_label = next((r[1] for r in DISMISSAL_REASONS if r[0] == reason), reason or "—")
    count = db.query(DismissalDoc).filter(DismissalDoc.doc_date >= doc_d.replace(day=1)).count()
    number = f"IB-{doc_d.strftime('%Y%m%d')}-{count + 1:04d}"
    doc = DismissalDoc(
        number=number,
        employee_id=emp.id,
        doc_date=doc_d,
        reason=reason_label,
        note=(note or "").strip() or None,
        user_id=current_user.id if current_user else None,
    )
    db.add(doc)
    db.flush()
    emp.is_active = False
    db.commit()
    return RedirectResponse(url=f"/dismissal/{doc.id}?created=1", status_code=303)


@router.get("/dismissal/{doc_id}", response_class=HTMLResponse)
async def dismissal_doc_view(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjati ko'rinishi."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee), joinedload(DismissalDoc.user))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    return templates.TemplateResponse("employees/dismissal_doc.html", {
        "request": request,
        "doc": doc,
        "current_user": current_user,
        "page_title": f"Ishdan bo'shatish {doc.number}",
    })


def _build_dismissal_docx(doc, company_name: str, employer_rep_name: str):
    """Ishdan bo'shatish hujjatini Word (.docx) sifatida qaytaradi (BytesIO)."""
    d = Document()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    d.add_heading("ISHDAN BO'SHATISH HAQIDA BUYRUQ", level=0)
    h = d.paragraphs[-1]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()
    p = d.add_paragraph()
    p.add_run(f"№ {doc.number}").bold = True
    p.add_run(f"   Sana: {doc.doc_date.strftime('%d.%m.%Y') if doc.doc_date else '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph()
    d.add_paragraph(f"Joy: ____________________________")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph()
    emp = doc.employee
    d.add_paragraph(
        f"1. {emp.full_name} (xodim kodi: {emp.code or '—'}), "
        f"{doc.doc_date.strftime('%d.%m.%Y')} sanadan boshlab ishdan bo'shatiladi."
    )
    d.add_paragraph(f"2. Ishdan bo'shatish sababi: {doc.reason or '—'}.")
    if doc.note:
        d.add_paragraph(f"3. Izoh: {doc.note}")
    d.add_paragraph()
    d.add_paragraph("Ish beruvchi:")
    d.add_paragraph(f"Korxona: {company_name}")
    d.add_paragraph(f"Rahbar: {employer_rep_name}")
    d.add_paragraph("Imzo: ______________________")
    d.add_paragraph()
    d.add_paragraph("Xodim bilan tanishtirildi:")
    d.add_paragraph(f"F.I.O: {emp.full_name}")
    d.add_paragraph("Imzo: ______________________")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@router.get("/dismissal/{doc_id}/export-word")
async def dismissal_doc_export_word(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_auth),
):
    """Ishdan bo'shatish hujjatini Word (.docx) formatida yuklab olish."""
    doc = (
        db.query(DismissalDoc)
        .options(joinedload(DismissalDoc.employee))
        .filter(DismissalDoc.id == doc_id)
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Hujjat topilmadi")
    company_name = "TOTLI HOLVA SWEETS"
    employer_rep_name = "Rahimov D.A."
    buf = _build_dismissal_docx(doc, company_name, employer_rep_name)
    safe_number = (doc.number or "ib").replace("/", "-").replace("\\", "-")
    filename = f"Ishdan_bo'shatish_{safe_number}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{quote(filename)}'},
    )
